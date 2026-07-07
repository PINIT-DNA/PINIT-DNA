"""
Generate PINIT-DNA_Global_Vision_Document.docx
Business / investor / enterprise presentation — not technical documentation.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1A, 0x2B, 0x4A)
TEAL = RGBColor(0x00, 0x7A, 0x8C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF8)
DARK_TEXT = RGBColor(0x22, 0x22, 0x22)

FONT = "Calibri"
FONT_TITLE = "Calibri Light"
BODY = Pt(11)
H1 = Pt(22)
H2 = Pt(16)
H3 = Pt(13)


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, bold=False, italic=False, size=None, color=None, font=None):
    run.font.name = font or FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font or FONT)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, bold=False, italic=False, align=None, size=None, color=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, size=size or BODY, color=color or DARK_TEXT)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run(run, bold=True, size={1: H1, 2: H2, 3: H3}.get(level, BODY), color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run, size=BODY, color=DARK_TEXT)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, italic=True, size=Pt(12), color=TEAL)


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=Pt(10), color=WHITE)
        shade_cell(cell, "1A2B4A")

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            text = row[c_idx] if c_idx < len(row) else ""
            run = cell.paragraphs[0].add_run(text)
            set_run(run, size=Pt(9), color=DARK_TEXT)
            if r_idx % 2 == 1:
                shade_cell(cell, "F4F6F8")
    doc.add_paragraph()
    return table


def add_toc_entry(doc, title: str, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(title)
    set_run(run, size=Pt(11 if level == 0 else 10), color=NAVY if level == 0 else DARK_TEXT, bold=(level == 0))


def add_static_toc(doc):
    add_heading(doc, "Table of Contents", level=1)
    entries = [
        "Introduction",
        "1. Executive Summary",
        "2. Current Global Problems",
        "3. Why Existing Solutions Are Not Enough",
        "4. What PINIT-DNA Provides",
        "5. User Journey — From Upload to Legal Proof",
        "6. Problem vs PINIT Solution",
        "7. Future Vision & Roadmap",
        "8. Target Users",
        "9. Why PINIT-DNA Is Different",
        "10. Long-Term Vision & Mission",
        "Appendix A — Industry Use Cases",
        "Appendix B — The Digital Ownership Gap",
        "Appendix C — Why Users Choose PINIT Over Cloud Storage",
    ]
    for entry in entries:
        add_toc_entry(doc, entry, level=0 if not entry.startswith("Appendix") else 0)
    doc.add_page_break()


def cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    add_para(
        doc,
        "PINIT-DNA",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(36),
        color=NAVY,
        space_after=12,
    )
    add_para(
        doc,
        "Digital Asset Identity, Protection & Investigation Platform",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(18),
        color=TEAL,
        space_after=24,
    )
    add_para(
        doc,
        "Protecting Digital Ownership for Every Individual and Organization",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(14),
        color=DARK_TEXT,
        space_after=36,
    )
    add_para(
        doc,
        "Business Vision & Global Opportunity Document",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(12),
        color=DARK_TEXT,
        space_after=48,
    )
    add_para(
        doc,
        f"July 2026  |  Confidential — For CEO, Investors & Enterprise Clients",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(10),
        color=RGBColor(0x88, 0x88, 0x88),
    )
    doc.add_page_break()

    add_heading(doc, "Introduction", level=1)
    intro = (
        "Every day, billions of digital files are created, shared, downloaded, edited, and reposted "
        "across the internet. Photos, videos, certificates, contracts, research, designs, and AI-generated "
        "content have become some of the most valuable assets people own — yet they have almost no legal "
        "identity, no ownership record, and no protection once they leave a device.\n\n"
        "Google Drive, Dropbox, WhatsApp, and Telegram were built to store and share files. They were "
        "not built to prove who created a file, detect when it was stolen, investigate suspicious copies, "
        "or produce court-ready evidence.\n\n"
        "PINIT-DNA exists to solve this gap. Our vision is to make enterprise-grade digital asset "
        "protection accessible to every individual — photographers, students, freelancers, influencers, "
        "lawyers, doctors, startups, and global enterprises alike.\n\n"
        "This document explains why the world needs PINIT-DNA, what problems it solves, and why users "
        "should protect their digital assets here — not only store them elsewhere."
    )
    add_para(doc, intro, space_after=12)
    doc.add_page_break()


# ── Data: Global Problems ─────────────────────────────────────────────────────

GLOBAL_PROBLEMS = [
    ("Image theft", "Photos downloaded from social media and reposted without credit", "Original creators lose revenue, reputation, and legal standing"),
    ("Video theft", "Short-form videos cropped and re-uploaded for monetization", "Creators lose ad revenue and brand value"),
    ("Certificate forgery", "Degrees and experience certificates edited in Photoshop", "Employers hire unqualified candidates; fraud goes undetected"),
    ("Document tampering", "Contracts and agreements modified after signing", "Legal disputes, financial loss, broken trust"),
    ("Resume editing", "Candidates inflate qualifications on shared PDFs", "Hiring risk and compliance failures"),
    ("Property document fraud", "Land papers and title deeds altered", "Property disputes and financial crime"),
    ("Medical report modification", "Scan reports and prescriptions edited", "Patient safety risk and insurance fraud"),
    ("AI-generated fake images", "Synthetic faces used for scams and impersonation", "Identity fraud and reputational damage"),
    ("Deepfake videos", "Public figures shown saying things they never said", "Political manipulation and personal harm"),
    ("Voice cloning", "AI voices used for fraud calls and blackmail", "Financial theft and emotional distress"),
    ("Screenshot misuse", "Confidential screens captured and shared widely", "Data leaks and privacy violations"),
    ("Screen recording leaks", "Private presentations recorded and distributed", "Trade secret exposure"),
    ("Fake invoices", "Bills edited to inflate or defraud payments", "Business and tax fraud"),
    ("Copyright theft", "Creative work copied without license", "Lost licensing income"),
    ("Logo theft", "Brand assets reused by competitors", "Brand dilution and customer confusion"),
    ("Research plagiarism", "Academic papers copied without attribution", "Integrity violations and IP loss"),
    ("Student assignment copying", "Projects submitted by multiple students", "Unfair grading and skill gaps"),
    ("Fake social media accounts", "Stolen photos used to build fake profiles", "Impersonation and scam victims"),
    ("Influencer content theft", "Reels and shorts reposted on fake pages", "Lost sponsorships and audience trust"),
    ("Course piracy", "Paid educational content redistributed free", "Creator revenue collapse"),
    ("Movie piracy", "Films leaked before release", "Studio revenue loss"),
    ("Music piracy", "Tracks ripped and re-uploaded", "Artist income destroyed"),
    ("Digital artwork theft", "NFT and digital art copied and resold", "Creator rights erased"),
    ("Government document forgery", "IDs and permits faked or altered", "National security and civic fraud"),
    ("Fake news", "Old images re-shared as breaking events", "Public misinformation at scale"),
    ("Evidence manipulation", "Crime scene or accident photos edited", "Justice system compromised"),
    ("Insurance fraud", "Accident photos and repair bills modified", "Millions lost by insurers annually"),
    ("Fake receipts", "Purchase proofs fabricated for refunds", "Retail and tax fraud"),
    ("Edited WhatsApp screenshots", "Chat conversations fabricated", "Personal and business defamation"),
    ("Court evidence modification", "Legal exhibits altered before submission", "Wrongful convictions or acquittals"),
    ("Identity theft", "Personal documents used to open accounts", "Financial and legal harm to victims"),
    ("Fake profile creation", "Someone's face used on dating or job sites", "Harassment and fraud"),
    ("Business confidential leaks", "Internal decks and strategy docs shared externally", "Competitive disadvantage"),
    ("Leaked company files", "Salary data and HR records exposed", "Employee trust destroyed"),
    ("Cloud storage misuse", "Shared links accessed by unauthorized parties", "Silent data exposure"),
    ("Data leakage", "Files forwarded beyond intended recipients", "Compliance breaches (GDPR, HIPAA)"),
    ("Internal employee leaks", "Staff exfiltrate files before leaving", "IP theft and sabotage"),
    ("Freelancer design theft", "Client refuses payment but keeps the logo", "Designer has no affordable proof"),
    ("Architecture drawing theft", "Blueprints copied by contractors", "Design IP stolen"),
    ("Patent theft", "Innovation documents copied before filing", "Startup destroyed before launch"),
    ("Startup idea theft", "Pitch decks forwarded to competitors", "First-mover advantage lost"),
    ("Source code theft", "Repositories cloned and republished", "Product copied overnight"),
    ("GitHub repository cloning", "Open or leaked repos mined for algorithms", "R&D investment wasted"),
    ("AI prompt theft", "Prompt libraries copied and monetized", "Creator competitive edge lost"),
    ("Model stealing", "ML weights and training data exfiltrated", "Years of research stolen"),
    ("Financial statement editing", "Balance sheets altered for loans", "Banking fraud"),
    ("Edited bank statements", "Transactions hidden or fabricated", "Loan and visa fraud"),
    ("Edited salary slips", "Income inflated for credit applications", "Credit risk for lenders"),
    ("Fake offer letters", "Employment letters created from templates", "Immigration and hiring fraud"),
    ("Edited ID cards", "Government IDs modified digitally", "KYC failures"),
    ("Edited passports", "Travel documents tampered for illegal entry", "Border security risk"),
    ("Edited Aadhaar / national ID", "Identity documents forged at scale", "Benefits and banking fraud"),
    ("Edited PAN / tax ID", "Tax identity documents falsified", "Revenue loss for governments"),
    ("Edited driving licenses", "License details changed after sharing", "Road safety and insurance fraud"),
    ("Wedding photography theft", "Bride's photos used by vendors without consent", "Emotional harm and unpaid usage"),
    ("Event photography misuse", "Event photos sold to third parties", "Photographer revenue stolen"),
    ("Family photo misuse", "Personal memories used in ads or scams", "Privacy violation"),
    ("Fake obituary images", "Deceased persons' photos misused online", "Family trauma"),
    ("Blackmail using edited media", "Fabricated images used for extortion", "Psychological and financial harm"),
    ("Cyber extortion", "Leaked private files used for ransom", "Personal safety threat"),
    ("Revenge editing", "Private media altered to harm reputation", "Irreversible social damage"),
    ("Social media impersonation", "Brands and people impersonated at scale", "Customer and follower deception"),
    ("WhatsApp forwarding chains", "Files lose origin after 50 forwards", "No accountability for content"),
    ("Telegram channel piracy", "Premium content reposted in piracy channels", "Subscription revenue lost"),
    ("LinkedIn portfolio theft", "Case studies copied by competitors", "Professional credibility stolen"),
    ("E-commerce product image theft", "Seller photos used by copycat listings", "Sales diverted to fakes"),
    ("Real estate listing fraud", "Property photos reused for fake listings", "Buyer scams"),
    ("Medical imaging fraud", "Scans swapped or edited for claims", "Healthcare system abuse"),
    ("Legal contract backdating", "Agreement dates changed after the fact", "Contract law undermined"),
    ("Academic certificate mills", "Fake degrees backed by doctored documents", "Credential inflation"),
    ("Journalist source leaks", "Unpublished materials forwarded prematurely", "Source safety compromised"),
    ("AI training data scraping", "Creators' work harvested without consent", "Creators unpaid for model value"),
]


# ── Data: Problem vs Solution (50+) ──────────────────────────────────────────

PROBLEM_SOLUTIONS = [
    ("Image theft on Instagram", "Anyone can download and repost; owner never knows", "Register image → DNA + Certificate → Monitor → Investigate → Evidence report", "Creators protect revenue and prove ownership"),
    ("YouTube Shorts re-upload", "Cropped clips monetized by strangers", "Video DNA + duplicate detection + future crawler alerts", "Creators reclaim ad revenue"),
    ("Fake degree certificate", "Photoshop edits pass HR checks", "DNA + Certificate ID + tamper verification", "Employers hire with confidence"),
    ("Contract edited after signing", "PDF text changed silently", "Original DNA hash vs suspect file comparison", "Legally defensible authenticity proof"),
    ("Resume inflation", "Shared PDF edited before interview", "Vault-stored original + verification link", "HR reduces hiring fraud"),
    ("Property paper fraud", "Title deeds altered", "Immutable DNA + ownership certificate + timeline", "Buyers verify before payment"),
    ("Medical report tampering", "Scan values changed for insurance", "Tamper detection + chain of custody", "Patient safety and insurer trust"),
    ("Deepfake CEO video", "Fake video authorizes wire transfer", "Future AI forensics + identity recovery", "Enterprises prevent wire fraud"),
    ("Voice clone scam call", "Family member's voice faked", "Future voice DNA registration", "Individuals verify authentic audio"),
    ("Confidential screenshot leak", "Board deck screenshotted and shared", "Screenshot detection + recipient tracking", "Enterprises trace leak source"),
    ("Screen recording of webinar", "Paid content recorded and resold", "Protected download + access intelligence", "Course creators protect revenue"),
    ("Fake invoice to finance", "Vendor bill amounts edited", "Certificate verification before payment", "AP teams stop payment fraud"),
    ("Logo copied by competitor", "Brand asset reused in ads", "DNA registration + duplicate detection", "Brand protection without lawyers first"),
    ("Research paper plagiarism", "PDF copied into another thesis", "DNA proves first creation date", "Universities uphold integrity"),
    ("Student assignment copy", "Same file submitted twice", "Duplicate upload detection inside platform", "Fair academic assessment"),
    ("Fake Instagram influencer page", "Stolen photos build fake following", "Crawler finds copies + identity recovery", "Influencers protect brand deals"),
    ("Paid course piracy", "Videos uploaded to Telegram", "Future internet monitoring + alerts", "Educators retain subscription income"),
    ("Film pre-release leak", "Copy spreads before box office", "Vault + tracking + investigation pipeline", "Studios reduce leak damage"),
    ("Music track rip", "Audio re-uploaded on streaming clones", "Audio DNA + future platform crawler", "Artists track unauthorized use"),
    ("Digital art NFT copy", "Artwork screenshot and resold", "Ownership certificate + global verification", "Collectors trust provenance"),
    ("Fake government ID", "National ID cards forged", "Verifier checks Certificate / DNA / Vault ID", "Institutions reduce KYC fraud"),
    ("Fake news image", "2010 photo shared as today's disaster", "Investigation recovers original + timeline", "Media verifies before publishing"),
    ("Crime scene photo edit", "Evidence altered before court", "Forensic DNA compare + chain of custody", "Justice system integrity"),
    ("Insurance accident fraud", "Damage photos staged or edited", "Original vault file vs claim photo", "Insurers save millions"),
    ("Fake purchase receipt", "Refund scams with edited receipts", "Merchant verifies certificate authenticity", "Retail loss prevention"),
    ("Fabricated WhatsApp chat", "Screenshots used in disputes", "Investigation + metadata analysis", "Parties prove truth in conflicts"),
    ("Court exhibit tampering", "Legal PDF modified", "Tamper verdict + evidence timeline", "Lawyers present admissible proof"),
    ("Stolen identity documents", "Passport scan used for accounts", "Protected vault + access logs", "Individuals control sensitive docs"),
    ("Dating app catfish", "Someone else's photos used", "Image investigation finds true owner", "Platforms reduce impersonation"),
    ("Startup pitch deck leak", "Investor forwards deck to rival", "Recipient tracking + watermark intelligence", "Founders protect ideas"),
    ("Employee exits with files", "Sales data taken on USB", "Vault access intelligence + audit trail", "HR and legal have evidence"),
    ("Freelancer unpaid logo", "Client uses design without paying", "Creation certificate + timeline", "Designer wins disputes faster"),
    ("Architect blueprint copy", "Contractor reuses design", "Document DNA + ownership proof", "Firm protects IP"),
    ("Patent race lost", "Idea documented but not protected", "Timestamped DNA + vault record", "Inventors establish priority"),
    ("GitHub code clone", "Repo copied to competing product", "Source code DNA + authorship proof", "Developers prove original work"),
    ("AI prompt marketplace theft", "Prompts copied and resold", "Creation history + certificate", "Prompt engineers monetize fairly"),
    ("ML model weights stolen", "Training artifacts exfiltrated", "Vault encryption + access forensics", "AI companies protect assets"),
    ("Edited bank statement", "Loan applicant hides debts", "Verifier confirms untampered original", "Banks reduce credit fraud"),
    ("Fake salary slip", "Income inflated for visa", "Employer-issued certificate verification", "Immigration integrity"),
    ("Fake job offer letter", "Candidate fabricates employment", "Certificate ID lookup shows fraud", "Recruiters filter faster"),
    ("Wedding photo commercial use", "Vendor sells couple's photos", "Monitoring finds unauthorized posts", "Couples control their memories"),
    ("Family photo in scam ad", "Grandparent's photo in fake charity", "Investigation traces original owner", "Families protect dignity"),
    ("Blackmail with edited media", "Fabricated compromising images", "Tamper analysis proves manipulation", "Victims have forensic defense"),
    ("Revenge porn editing", "Private media altered and shared", "Evidence report + chain of custody", "Legal action supported"),
    ("Brand impersonation page", "Logo + content copied", "Crawler + duplicate alerts", "Brands protect customers"),
    ("E-commerce image theft", "Product photos on fake stores", "Internet monitoring finds copies", "Sellers protect listings"),
    ("Real estate scam listing", "Stolen photos for fake property", "Verify vault owner before deposit", "Buyers avoid fraud"),
    ("Doctor prescription edit", "Dosage changed on shared PDF", "Tamper detection on medical docs", "Pharmacies verify authenticity"),
    ("Lawyer evidence package", "Client sends altered contract", "Unified investigation in seconds", "Firms work faster with proof"),
    ("Journalist image verify", "Viral image needs source check", "Recover original + edit history", "Newsrooms reduce misinformation"),
    ("AI-generated fake certificate", "Synthetic document passes glance test", "DNA + certificate not in registry = fraud", "Institutions stay ahead of AI fraud"),
    ("Influencer reel repost", "Brand content on fake account", "Automatic alert when copy found", "Influencers protect sponsorships"),
    ("Musician leaked demo", "Unreleased track shared early", "Audio DNA + monitoring", "Labels control release timing"),
    ("University thesis copy", "Graduate work plagiarized", "First-upload DNA proves author", "Graduates protect years of work"),
    ("Corporate M&A data room leak", "Due diligence files forwarded", "Recipient intelligence + access logs", "Deal teams contain breaches"),
    ("Government tender fraud", "Bid documents altered", "Chain of custody + tamper check", "Public procurement integrity"),
    ("Telegram piracy channel", "Premium PDFs redistributed", "Future crawler on messaging surfaces", "Publishers recover revenue"),
    ("Cloud link oversharing", "Drive link indexed by Google", "Controlled sharing via PINIT links", "Users know who accessed what"),
    ("Years-later ownership dispute", "Creator cannot prove 'I made this first'", "Permanent DNA + vault + certificate + timeline", "Time-proof ownership record"),
    ("Suspicious file from unknown sender", "No tool to check origin", "Unified Investigation Center auto-recovers owner", "Anyone can verify before trusting"),
    ("No file history after sharing", "Once sent, visibility is zero", "Evidence timeline: share, download, investigate", "Full lifecycle transparency"),
    ("Enterprise-only forensics cost", "DRM tools cost lakhs per year", "Same-grade protection for individuals", "Democratized digital rights"),
    ("Antivirus misses file theft", "AV protects device, not asset identity", "PINIT protects the file's identity everywhere", "Ownership layer beyond antivirus"),
    ("Cloud stores but doesn't prove", "Drive holds bytes, not proof", "Vault + DNA + certificate = identity", "Storage plus legal standing"),
]


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY

    # ── Cover ─────────────────────────────────────────────────────────────────
    cover_page(doc)
    add_static_toc(doc)

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        "Today, every person and organization owns valuable digital assets. These include photographs, "
        "videos, certificates, legal contracts, medical records, research papers, music, software source "
        "code, architectural drawings, financial statements, and AI-generated content. The total value of "
        "these assets often exceeds the value of physical possessions — yet they have no registration, no "
        "ownership record, and no protection once they leave a phone or laptop.",
    )
    add_para(
        doc,
        "When someone downloads your photo, edits your certificate, copies your video, or forwards your "
        "confidential document, you typically discover it only after the damage is done. Google Drive, "
        "Dropbox, OneDrive, WhatsApp, and Telegram solve storage and sharing. They do not solve ownership, "
        "authenticity, investigation, or legal proof.",
    )
    add_para(
        doc,
        "PINIT-DNA is a Digital Asset Identity, Protection & Investigation Platform. It gives every file a "
        "permanent Digital DNA identity, stores it in an encrypted vault, issues an ownership certificate, "
        "monitors for unauthorized use, investigates suspicious copies, and produces forensic evidence "
        "reports suitable for copyright claims, legal disputes, and enterprise compliance.",
        bold=False,
    )
    add_quote(
        doc,
        "Our vision is to make enterprise-grade digital asset protection accessible to every individual — "
        "not only to Disney, Netflix, and Fortune 500 security teams.",
    )
    doc.add_page_break()

    # ── 2. Current Global Problems ────────────────────────────────────────────
    add_heading(doc, "2. Current Global Problems", level=1)
    add_para(
        doc,
        "The following table maps real-world digital threats to their current situation and business or "
        "personal impact. These are not edge cases — they affect millions of creators, professionals, "
        "students, and organizations every day.",
    )
    add_table(
        doc,
        ["Problem", "Current Situation", "Impact"],
        [[a, b, c] for a, b, c in GLOBAL_PROBLEMS],
    )
    doc.add_page_break()

    # ── 3. Why Existing Solutions Are Not Enough ──────────────────────────────
    add_heading(doc, "3. Why Existing Solutions Are Not Enough", level=1)
    add_para(
        doc,
        "Mainstream tools were designed for convenience — not for proving who owns a file, whether it was "
        "tampered with, or where unauthorized copies appear online.",
    )
    add_table(
        doc,
        ["Platform / Tool", "What It Provides", "What It Does NOT Provide"],
        [
            ("Google Drive", "Storage, sharing, sync, basic permissions", "DNA identity, tamper detection, forensics, ownership certificate, crawler, evidence timeline"),
            ("Dropbox", "Cloud backup, file sync, link sharing", "Investigation, identity recovery, leak intelligence, legal evidence reports"),
            ("OneDrive", "Microsoft ecosystem storage", "Digital ownership layer, duplicate theft detection, certificate verification"),
            ("WhatsApp", "Fast messaging and file forwarding", "Ownership proof, tamper alerts, access forensics, controlled revocation"),
            ("Telegram", "Large file sharing, channels", "File identity, monitoring, investigation, chain of custody"),
            ("iCloud", "Apple device backup and photos", "Cross-platform verification, forensic comparison, internet leak detection"),
            ("Local hard drive / USB", "Offline storage", "Encryption with identity, sharing intelligence, remote monitoring"),
            ("Enterprise DRM only", "Studios and large corporations", "Affordable access for individuals, freelancers, and SMBs"),
        ],
    )
    add_para(doc, "Summary comparison of capability categories:", bold=True)
    add_table(
        doc,
        ["Capability", "Drive / Dropbox / WhatsApp", "PINIT-DNA"],
        [
            ("File storage", "Yes", "Yes (encrypted Vault)"),
            ("File sharing", "Yes", "Yes (Smart Links with intelligence)"),
            ("Backup & sync", "Yes", "Vault-backed secure storage"),
            ("Ownership identity (DNA)", "No", "Yes — every file"),
            ("Ownership certificate", "No", "Yes — verifiable globally"),
            ("Tamper detection", "No", "Yes"),
            ("Unified investigation", "No", "Yes"),
            ("Identity recovery", "No", "Yes"),
            ("Internet crawler / monitoring", "No", "Yes (roadmap: full automation)"),
            ("Evidence timeline", "No", "Yes"),
            ("Forensic evidence report", "No", "Yes"),
            ("Recipient & access intelligence", "No", "Yes"),
            ("Leak detection", "No", "Yes"),
        ],
    )
    doc.add_page_break()

    # ── 4. What PINIT-DNA Provides ────────────────────────────────────────────
    add_heading(doc, "4. What PINIT-DNA Provides", level=1)
    add_table(
        doc,
        ["Feature", "Purpose", "Benefit to User"],
        [
            ("Encrypted Vault", "Secure storage with tenant isolation", "Files protected at rest; only owner controls access"),
            ("Digital DNA (15-layer fingerprint)", "Unique immutable identity per file", "Survives crop, compress, screenshot — still identifiable"),
            ("Ownership Certificate", "Official proof of registration", "Shareable proof for clients, courts, and employers"),
            ("Smart Share Links", "Controlled sharing with tracking", "Know who opened, downloaded, forwarded"),
            ("Access Intelligence", "GPS, device, and session analytics", "See exactly how files are accessed"),
            ("Recipient Tracking", "Watermark and recipient profiles", "Trace which recipient leaked content"),
            ("Monitoring", "Enroll assets for ongoing watch", "Stay alerted when files appear online"),
            ("Internet Crawler", "Search web for unauthorized copies", "Find theft you would never discover manually"),
            ("Duplicate Detection", "Block and alert on copy uploads", "Prevent others from claiming your work"),
            ("Unified Investigation Center", "Upload suspect file → auto-investigate", "Recover owner, vault, certificate in seconds"),
            ("Identity Recovery", "Match unknown file to original owner", "Answer 'who really owns this?'"),
            ("Tamper Detection", "Compare suspect vs vault original", "Prove editing, cropping, or forgery"),
            ("Evidence Timeline", "Append-only lifecycle events", "Court-ready chain of custody"),
            ("Forensic Reports", "PDF evidence packages", "Copyright claims, legal disputes, insurance"),
            ("Certificate Verification", "Anyone verifies via ID lookup", "Global trust without contacting owner first"),
            ("Leak Intelligence", "Attribute leaked copies to source", "Find how and where content escaped"),
            ("Protected Download (TEP)", "Tracked, revocable downloads", "Control files even after sending"),
            ("Semantic AI Search", "Find files by meaning in vault", "Enterprise-grade discovery for everyone"),
            ("Future: AI Forensics", "Deepfake and synthetic media detection", "Trust engine for the AI era"),
            ("Future: Auto Investigation", "Crawler match triggers full case", "Zero manual work for owners"),
            ("Future: Global Monitoring", "Social platform coverage at scale", "Instagram, YouTube, TikTok, Telegram alerts"),
        ],
    )
    doc.add_page_break()

    # ── 5. User Journey ───────────────────────────────────────────────────────
    add_heading(doc, "5. User Journey — From Upload to Legal Proof", level=1)
    add_para(
        doc,
        "The following flow shows how a normal user — a photographer, student, lawyer, or business — "
        "protects a file from creation through to forensic evidence.",
    )
    flow_steps = [
        "User uploads file (photo, video, PDF, contract, design, code)",
        "↓",
        "PINIT generates unique Digital DNA (15-layer forensic fingerprint)",
        "↓",
        "File encrypted with AES-256-GCM and stored in secure Vault",
        "↓",
        "Ownership Certificate issued with verifiable Certificate ID",
        "↓",
        "Monitoring automatically enrolled (optional watch URLs)",
        "↓",
        "Internet Crawler searches for unauthorized copies (roadmap: fully automatic)",
        "↓",
        "If match found → Owner receives real-time alert notification",
        "↓",
        "Automatic Investigation starts (Unified Investigation Center)",
        "↓",
        "System recovers original owner, vault record, and certificate",
        "↓",
        "Tamper Detection compares suspect file vs original",
        "↓",
        "Evidence Timeline and Forensic Report generated",
        "↓",
        "Owner has legal proof for copyright claim, court, or internal audit",
    ]
    for step in flow_steps:
        if step == "↓":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("▼")
            set_run(run, color=TEAL, size=Pt(10))
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(step)
            set_run(run, size=Pt(11), color=NAVY, bold=("upload" in step.lower() or "legal proof" in step.lower()))

    add_para(doc, "", space_after=6)
    add_heading(doc, "Future State — Fully Automatic Protection", level=2)
    add_para(
        doc,
        "Imagine uploading one photo today. Five years later, someone posts a cropped version on Instagram. "
        "The crawler finds it. PINIT sends: 'Your protected image has appeared online.' You tap once. "
        "Within seconds you see: original owner, first upload date, similarity score, tampering analysis, "
        "sharing history, certificate, and a downloadable evidence report. No manual investigation required.",
    )
    doc.add_page_break()

    # ── 6. Problem vs PINIT Solution ──────────────────────────────────────────
    add_heading(doc, "6. Problem vs PINIT Solution", level=1)
    add_para(doc, "Fifty-plus real-world scenarios and how PINIT-DNA addresses each one.", space_after=8)
    add_table(
        doc,
        ["Problem", "Current Situation", "PINIT Solution", "Business Value"],
        PROBLEM_SOLUTIONS,
    )
    doc.add_page_break()

    # ── 7. Future Vision ──────────────────────────────────────────────────────
    add_heading(doc, "7. Future Vision & Roadmap", level=1)
    add_para(
        doc,
        "PINIT-DNA is building toward a world where every digital asset has identity, every unauthorized "
        "use is detectable, and every individual has access to forensic-grade protection.",
    )
    roadmap = [
        ("Automatic Internet Crawler", "Continuously scan web and social platforms for leaked copies"),
        ("AI-Powered Investigation", "Machine learning ranks matches and auto-builds case files"),
        ("Social Media Monitoring", "Instagram, Facebook, YouTube, TikTok, Pinterest detection"),
        ("Messaging Surface Monitoring", "Telegram and public channel leak detection"),
        ("Dark Web Monitoring", "Early warning for stolen corporate and personal assets"),
        ("OCR & Document Intelligence", "Extract and verify text in certificates and contracts"),
        ("AI-Generated Image Detection", "Identify synthetic vs authentic media"),
        ("Deepfake & Voice DNA", "Register and verify voice and video authenticity"),
        ("Multi-Media DNA", "Unified identity across image, video, audio, document, code"),
        ("Blockchain Anchoring", "Optional immutable timestamp on public ledger"),
        ("Global Certificate Verification", "Anyone worldwide verifies ownership in one click"),
        ("Real-Time Leak Alerts", "Push notifications the moment a copy is found"),
        ("Automatic Case Creation", "Crawler hit opens full investigation automatically"),
        ("Court-Ready Report Generator", "One-click PDF for legal and compliance teams"),
        ("AI Evidence Collection", "Auto-gather screenshots, URLs, and metadata"),
        ("Global Ownership Registry", "Searchable trust layer for digital assets"),
        ("Recipient Behavior Intelligence", "Pattern analysis on who leaks content"),
        ("Zero-Knowledge Proof Options", "Verify ownership without exposing file content"),
        ("Mobile App", "Protect assets from phone camera in real time"),
        ("Browser Extension", "Right-click → Protect with PINIT"),
        ("Enterprise Dashboard", "Organization-wide asset governance"),
        ("Government & Institutional Integration", "APIs for universities, courts, and agencies"),
        ("Developer API Platform", "Embed PINIT protection in any product"),
    ]
    add_table(doc, ["Future Capability", "Description"], roadmap)
    doc.add_page_break()

    # ── 8. Target Users ───────────────────────────────────────────────────────
    add_heading(doc, "8. Target Users", level=1)
    add_table(
        doc,
        ["Segment", "Primary Pain", "How PINIT Helps"],
        [
            ("Individual users", "Personal photos and documents stolen or misused", "Digital locker with identity and proof"),
            ("Influencers & creators", "Reels, shorts, and brand content reposted", "Monitoring, alerts, and copyright evidence"),
            ("YouTubers & videographers", "Video theft and monetization loss", "Video DNA and duplicate detection"),
            ("Photographers", "Wedding and commercial image theft", "Ownership certificate and leak tracking"),
            ("Lawyers & legal firms", "Contract tampering and evidence disputes", "Tamper detection and chain of custody"),
            ("Doctors & clinics", "Medical report modification", "Verify originals before treatment decisions"),
            ("Students & researchers", "Plagiarism and idea theft", "Prove first creation with timestamped DNA"),
            ("Freelancers & designers", "Clients use work without payment", "Creation proof and timeline"),
            ("Architects & engineers", "Blueprint and CAD theft", "IP protection for design files"),
            ("Startups & businesses", "Pitch decks and IP leaks", "Recipient tracking and access intelligence"),
            ("Schools & universities", "Academic integrity", "Assignment and thesis verification"),
            ("Government agencies", "Document forgery and fraud", "Certificate verification at scale"),
            ("Police & forensic labs", "Digital evidence analysis", "Investigation center and reports"),
            ("Insurance companies", "Claim photo and invoice fraud", "Compare claim media vs originals"),
            ("Media & news agencies", "Fake news image verification", "Recover original source fast"),
            ("Banks & financial firms", "Edited statements and KYC fraud", "Verify untampered documents"),
            ("Musicians & labels", "Track leaks and unauthorized uploads", "Audio DNA and monitoring"),
            ("Software developers", "Code and algorithm theft", "Source authorship proof"),
            ("AI companies", "Model and prompt IP protection", "Vault + access forensics"),
            ("Corporate enterprises", "Internal leaks and compliance", "Enterprise dashboard and audit trails"),
        ],
    )
    doc.add_page_break()

    # ── 9. Why PINIT Is Different ───────────────────────────────────────────────
    add_heading(doc, "9. Why PINIT-DNA Is Different", level=1)
    add_para(
        doc,
        "Large enterprises already spend millions on Digital Rights Management (DRM), forensic watermarking, "
        "data loss prevention, and cybersecurity suites. Disney, Netflix, Adobe, and global banks have teams "
        "of analysts, lawyers, and forensic vendors. Ordinary people — photographers, students, freelancers, "
        "influencers, and small businesses — have almost nothing affordable.",
    )
    add_para(
        doc,
        "PINIT-DNA closes this gap. We are not building another cloud drive. We are building the ownership "
        "layer the internet never had — combining vault storage, forensic DNA identity, certificates, "
        "monitoring, investigation, and evidence into one platform designed for everyone.",
    )
    add_heading(doc, "What makes PINIT unique", level=2)
    for item in [
        "Enterprise-grade forensic DNA — accessible to individuals, not only corporations",
        "End-to-end journey: protect → store → share → monitor → investigate → prove",
        "Works across images, videos, documents, audio, archives, and more",
        "Verification anyone can use — Certificate ID, DNA ID, or Vault ID lookup",
        "Real-time notifications when your content is accessed or at risk",
        "Evidence timeline and chain of custody — not just file bytes in a folder",
        "Vision for automatic internet monitoring — your copyright assistant",
        "Built for the AI era — when nothing on screen can be trusted without proof",
    ]:
        add_bullet(doc, item)
    add_quote(
        doc,
        "Google Drive answers: 'Where is my file?'  —  PINIT-DNA answers: 'Who owns it, was it tampered with, "
        "where else does it exist, and can I prove it in court?'",
    )
    doc.add_page_break()

    # ── 10. Long-Term Vision ──────────────────────────────────────────────────
    add_heading(doc, "10. Long-Term Vision & Mission", level=1)
    visions = [
        ("Today, antivirus protects computers.", "Tomorrow, PINIT protects digital ownership."),
        ("Today, cloud storage stores files.", "Tomorrow, PINIT protects the identity of every file."),
        ("Today, cybersecurity protects systems and networks.", "Tomorrow, PINIT protects every digital asset people create."),
        ("Today, only enterprises can afford digital forensics.", "Tomorrow, every individual has the same level of protection."),
    ]
    for today, tomorrow in visions:
        add_para(doc, today, bold=True, color=NAVY)
        add_para(doc, tomorrow, color=TEAL, space_after=12)

    add_para(doc, "", space_after=6)
    add_heading(doc, "Mission Statement", level=2)
    add_quote(
        doc,
        "PINIT-DNA exists to give every person and every organization a permanent digital identity for "
        "their most valuable files — so that ownership can be proven, misuse can be detected, and truth "
        "can survive in a world of endless copies, edits, and AI-generated fakes.",
    )
    add_para(
        doc,
        "We believe the internet needs an ownership layer — the digital equivalent of a vehicle registration "
        "or land title — for photos, videos, certificates, contracts, code, and every asset that matters. "
        "PINIT-DNA is building that layer for the world.",
        space_after=24,
    )

    doc.add_page_break()

    # ── Appendix A: Industry Use Cases ────────────────────────────────────────
    add_heading(doc, "Appendix A — Industry Use Cases", level=1)
    use_cases = [
        (
            "Photographers & Wedding Studios",
            "A wedding photographer delivers 500 edited photos to the couple. Within weeks, a decorator "
            "reposts images on Instagram without credit. A stock site sells them. The photographer has "
            "no affordable way to prove ownership or find every copy. With PINIT-DNA, each deliverable is "
            "registered at upload, certificates are issued, and monitoring alerts the studio when copies "
            "appear online — with a forensic report ready for a copyright notice.",
        ),
        (
            "Influencers & Content Creators",
            "A creator spends hours on a reel. Someone downloads it, crops the watermark, and reposts on "
            "a fake page that earns brand deals. The creator loses sponsorship revenue and reputation. "
            "PINIT-DNA registers the original, tracks shares, and future crawler automation finds "
            "unauthorized reposts — turning PINIT into a personal copyright assistant.",
        ),
        (
            "Freelancers & Designers",
            "A designer delivers a logo. The client refuses final payment but uses the logo on their "
            "website. Legal action costs more than the project fee. PINIT-DNA provides a timestamped "
            "creation record, ownership certificate, and evidence timeline — affordable proof that "
            "levels the playing field.",
        ),
        (
            "Lawyers & Legal Professionals",
            "A law firm receives a contract PDF from opposing counsel. Was it altered after signing? "
            "Today, verification requires expensive forensic vendors. PINIT-DNA lets any party upload "
            "the suspect file and instantly compare it against the vault original — tamper verdict, "
            "chain of custody, and downloadable evidence report included.",
        ),
        (
            "Students & Universities",
            "A graduate student's thesis is copied and submitted by another student two semesters later. "
            "Proving who created first is nearly impossible. PINIT-DNA's DNA registration timestamps "
            "the original at creation — giving universities a fair, automated integrity tool.",
        ),
        (
            "Insurance Companies",
            "Claimants submit edited accident photos or inflated repair invoices. Adjusters cannot "
            "easily verify authenticity. PINIT-DNA lets insurers request certificate-verified originals "
            "and compare submitted media against vault records — reducing fraud at scale.",
        ),
        (
            "Healthcare Providers",
            "Medical scans and prescriptions shared via WhatsApp can be edited before reaching a "
            "pharmacy or second opinion. PINIT-DNA vault storage with tamper detection ensures "
            "patients and providers trust the document in front of them.",
        ),
        (
            "Startups & Enterprises",
            "A pitch deck forwarded from an investor data room appears on a competitor's desk. "
            "Recipient tracking and access intelligence show exactly who opened, downloaded, and "
            "forwarded the file — containing leaks before they destroy a funding round.",
        ),
    ]
    for title, body in use_cases:
        add_heading(doc, title, level=2)
        add_para(doc, body, space_after=10)
    doc.add_page_break()

    # ── Appendix B: Digital Ownership Gap ─────────────────────────────────────
    add_heading(doc, "Appendix B — The Digital Ownership Gap", level=1)
    add_para(
        doc,
        "Every valuable physical asset in the modern world has an identity system. Cars have registration "
        "numbers and chassis IDs. Land has survey numbers and title deeds. Houses have ownership records "
        "in government registries. Yet a digital photograph worth lakhs in licensing revenue — or a legal "
        "contract worth crores — has no equivalent registration when it leaves your phone.",
    )
    add_para(
        doc,
        "This is the gap PINIT-DNA fills. We are building the registration layer for digital assets — "
        "not to replace Google Drive or Dropbox, but to sit above storage and answer the questions "
        "cloud platforms were never designed to answer:",
    )
    for q in [
        "Who originally created this file?",
        "Has it been edited, cropped, or forged?",
        "Where else does this file exist on the internet?",
        "Can I prove ownership in a copyright claim or court?",
        "Who accessed or leaked this file after I shared it?",
    ]:
        add_bullet(doc, q)
    add_quote(
        doc,
        "Storage answers WHERE your file is. PINIT-DNA answers WHO owns it, WHETHER it is authentic, "
        "and WHAT happened to it after you shared it.",
    )
    doc.add_page_break()

    # ── Appendix C: Why Choose PINIT ──────────────────────────────────────────
    add_heading(doc, "Appendix C — Why Users Choose PINIT Over Cloud Storage Alone", level=1)
    add_table(
        doc,
        ["User Question", "Google Drive / WhatsApp", "PINIT-DNA"],
        [
            ("Where is my file?", "Yes — stored in cloud or chat", "Yes — encrypted Vault"),
            ("Who owns this file?", "No ownership record", "DNA + Certificate"),
            ("Was this file edited?", "No detection", "Tamper analysis vs original"),
            ("Who downloaded or forwarded it?", "Limited or no logs", "Full access intelligence"),
            ("Is this copy stolen from me?", "No way to know", "Monitoring + crawler + investigation"),
            ("Can I prove this in court?", "No forensic report", "Evidence timeline + PDF report"),
            ("Can anyone verify authenticity?", "No public verification", "Certificate / DNA / Vault ID lookup"),
            ("What if someone screenshots my confidential file?", "No protection", "Screenshot detection + recipient tracking"),
            ("What happens in 5 years if someone reposts my photo?", "You will never know", "Crawler alerts + auto-investigation (roadmap)"),
        ],
    )
    add_para(
        doc,
        "People do not switch to PINIT because they need another folder. They switch because they are "
        "tired of discovering theft too late, tired of having no proof, and tired of enterprise-grade "
        "protection being reserved for companies that can afford million-dollar security budgets.",
        space_after=12,
    )
    add_para(
        doc,
        "— PINIT-DNA  |  Digital Asset Identity, Protection & Investigation Platform",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(10),
        color=RGBColor(0x88, 0x88, 0x88),
    )

    return doc


def main():
    out = Path(__file__).resolve().parent / "PINIT-DNA_Global_Vision_Document.docx"
    doc = build_document()
    try:
        doc.save(str(out))
    except PermissionError:
        out = Path(__file__).resolve().parent / "PINIT-DNA_Global_Vision_Document_v2.docx"
        doc.save(str(out))
        print("Original file is open — saved as v2 instead.")
    print(f"Created: {out}")


if __name__ == "__main__":
    main()
