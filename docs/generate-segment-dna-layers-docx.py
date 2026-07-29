"""
Generate docs/PINIT-DNA_Segment_15_Layer_Code_Capability.docx

World-facing PINIT presentation for ~40 creator/buyer segments.
Per segment:
  1) Responsibilities
  2) Why they use PINIT
  3) How PINIT helps
  4) 15-layer table — Yes/No + how DNA generates for THEIR assets
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1A, 0x2B, 0x4A)
TEAL = RGBColor(0x00, 0x7A, 0x8C)
DARK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"
BODY = Pt(10)
SMALL = Pt(9)
H1 = Pt(22)
H2 = Pt(16)
H3 = Pt(13)

LAYERS = [
    (1, "Cryptographic", "Exact byte identity"),
    (2, "Structural", "Layout / container structure"),
    (3, "Perceptual", "Survives crop / compress / re-encode"),
    (4, "Semantic", "Content look / meaning fingerprint"),
    (5, "Metadata", "EXIF / file provenance hash"),
    (6, "Signature", "Integrity seal (LSB / HMAC)"),
    (7, "Behavioral", "Who uploaded / session fingerprint"),
    (8, "Relationship", "Links related vault assets"),
    (9, "Origin", "Upload IP / device / geo context"),
    (10, "Evolution", "Version history chain"),
    (11, "Deepfake Detection", "AI / manipulation signal"),
    (12, "Invisible Watermark", "Owner watermark record"),
    (13, "Chain of Custody", "Legal custody entry"),
    (14, "ZK Ownership Proof", "Ownership commitment proof"),
    (15, "Biometric Bind", "Bind to owner biometrics"),
]

# segment: role, why_pinit, how_helps, assets, asset_examples (list of (file, engine)), no_case
# rows built dynamically from assets + media_kind

Segment = dict  # typed loosely for clarity in data


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run(run, *, bold=False, size=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color


def add_para(doc, text, *, bold=False, size=None, color=None, space_after=4, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run(run, bold=bold, size=size or BODY, color=color or DARK)
    return p


def add_label_value(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}: ")
    set_run(r1, bold=True, size=SMALL, color=NAVY)
    r2 = p.add_run(value)
    set_run(r2, size=SMALL, color=DARK)


CREATORS: list[dict] = [
    {
        "name": "Photographers",
        "role": "Shoot, edit, and deliver stills for clients, stock, weddings, and brands; own visual copyright.",
        "why": "Final photos are downloaded from social, stolen from portfolios, and resold without credit.",
        "help": "Register finals before publish → ownership certificate → monitor leaks → investigate with forensic DNA match.",
        "assets": "IMAGE (JPEG/PNG/TIFF/WebP); optional PDF contact sheets",
        "examples": [("Wedding/commercial JPEG", "Image DNA"), ("Portfolio PNG", "Image DNA")],
        "no": "Camera RAW (.CR2/.NEF) — export JPEG/TIFF first.",
        "media": "image",
    },
    {
        "name": "Videographers",
        "role": "Produce and deliver video masters, highlight reels, and client films.",
        "why": "Reels and masters are re-uploaded to YouTube/social after re-encode or crop.",
        "help": "DNA the master + key stills before platform upload; protect-download for client previews; investigate pirated cuts.",
        "assets": "VIDEO (MP4/MOV); IMAGE stills",
        "examples": [("Master MP4", "Video DNA"), ("Thumbnail JPEG", "Image DNA")],
        "no": "NLE projects (.prproj/.aep) — export MP4 first.",
        "media": "video",
    },
    {
        "name": "Graphic Designers",
        "role": "Create brand kits, posters, packaging, and client visual systems.",
        "why": "Deliverable exports get reused by clients beyond license or scraped online.",
        "help": "DNA approved PNG/PDF/PPTX finals; sealed client handoff; proof when work is reused without rights.",
        "assets": "IMAGE, PDF, PPTX",
        "examples": [("Poster PNG", "Image DNA"), ("Brand PDF", "PDF DNA"), ("Deck PPTX", "PPTX DNA")],
        "no": "PSD/AI/Figma native — export PNG/PDF/PPTX first.",
        "media": "mixed",
    },
    {
        "name": "UI/UX Designers",
        "role": "Design product interfaces, flows, and stakeholder presentation decks.",
        "why": "Unreleased UI screens and pitch decks leak to competitors or social.",
        "help": "DNA shipped mockups and decks; Smart Link for stakeholder review; ownership proof for product visual IP.",
        "assets": "IMAGE (PNG), PDF, PPTX",
        "examples": [("Screen PNG", "Image DNA"), ("Spec PDF", "PDF DNA")],
        "no": "Figma/Sketch native — export PNG/PDF first.",
        "media": "image",
    },
    {
        "name": "Motion Designers",
        "role": "Design and deliver animated spots, explainers, and social motion.",
        "why": "Rendered finals are ripped from social and reused in other ads.",
        "help": "DNA rendered MP4 + still frames; watermark client previews; match re-encoded social copies.",
        "assets": "VIDEO renders; IMAGE frames",
        "examples": [("Spot MP4", "Video DNA"), ("Frame PNG", "Image DNA")],
        "no": "AE/Cinema projects — export MP4/PNG first.",
        "media": "video",
    },
    {
        "name": "3D Artists",
        "role": "Create models, scenes, and marketplace-ready beauty renders.",
        "why": "Renders and turntables are scraped from ArtStation/marketplaces and resold.",
        "help": "DNA beauty renders and turntable videos before posting; ownership proof for marketplace disputes.",
        "assets": "IMAGE renders; VIDEO turntables",
        "examples": [("Beauty PNG", "Image DNA"), ("Turntable MP4", "Video DNA")],
        "no": ".blend/.fbx/.obj — export IMAGE/VIDEO first.",
        "media": "image",
    },
    {
        "name": "Architects",
        "role": "Issue drawings, visualizations, and design packages to clients and contractors.",
        "why": "Renders and plan PDFs are stolen for fake listings or unauthorized builds.",
        "help": "DNA issued PDFs and marketing renders; custody on issued sets; prove ownership in IP disputes.",
        "assets": "PDF plans; IMAGE renders; optional PPTX",
        "examples": [("Plan PDF", "PDF DNA"), ("Viz JPEG", "Image DNA")],
        "no": "DWG/RVT/IFC — export PDF/PNG first.",
        "media": "doc",
    },
    {
        "name": "Musicians",
        "role": "Create, mix, and release masters and cover art.",
        "why": "Tracks and artwork are ripped and rehosted without credit or royalties.",
        "help": "DNA masters + covers before DistroKid/social; watermark private listens; ownership for copyright claims.",
        "assets": "AUDIO (WAV/FLAC/MP3); IMAGE covers",
        "examples": [("Master WAV", "Audio DNA"), ("Cover JPG", "Image DNA")],
        "no": "DAW projects — bounce audio first.",
        "media": "audio",
    },
    {
        "name": "Film Studios",
        "role": "Produce, post, and distribute film/TV masters, stills, and scripts under rights control.",
        "why": "Screeners, stills, and scripts leak to pirate networks before release.",
        "help": "Full DNA on masters/scripts; TEP/Smart Link screeners; custody + watermark + investigate piracy.",
        "assets": "VIDEO, IMAGE, PDF, DOCX, ZIP",
        "examples": [("Master MP4", "Video DNA"), ("Script PDF", "PDF DNA"), ("Press ZIP", "ZIP DNA")],
        "no": "Proprietary camera wrappers — mezzanine MP4 first if needed.",
        "media": "video",
    },
    {
        "name": "Podcasters",
        "role": "Record, edit, and publish episodes with show art.",
        "why": "Paid/early episodes and covers get rehosted or ripped from RSS/YouTube.",
        "help": "DNA episode + art before publish; protect early-access downloads; match ripped rehosts.",
        "assets": "AUDIO; IMAGE; optional TXT transcript",
        "examples": [("Episode MP3", "Audio DNA"), ("Cover PNG", "Image DNA")],
        "no": "None for standard episode+cover LIVE files.",
        "media": "audio",
    },
    {
        "name": "Writers & Authors",
        "role": "Write, revise, and submit manuscripts, proofs, and covers.",
        "why": "Drafts and proofs leak online; plagiarism and unauthorized distribution risk.",
        "help": "DNA manuscript versions; watermark review PDFs; ownership proof for disputes.",
        "assets": "DOCX, PDF, TXT; IMAGE covers",
        "examples": [("Manuscript DOCX", "DOCX DNA"), ("Proof PDF", "PDF DNA")],
        "no": "None for standard manuscript LIVE types.",
        "media": "doc",
    },
    {
        "name": "Journalists",
        "role": "Capture exclusives, handle sources, and publish under embargo rules.",
        "why": "Exclusives and source docs can leak before publish or be stolen.",
        "help": "DNA exclusives; Smart Link for embargo review; custody on source files; investigate early leaks.",
        "assets": "IMAGE, VIDEO, PDF, DOCX",
        "examples": [("Exclusive photo", "Image DNA"), ("Source PDF", "PDF DNA")],
        "no": "None for typical exclusive LIVE files.",
        "media": "mixed",
    },
    {
        "name": "Educators",
        "role": "Build and sell or deliver courses, slides, and lesson videos.",
        "why": "Paid course packs are pirated and redistributed for free.",
        "help": "DNA curriculum packs; watermark student downloads; monitor pirated re-uploads.",
        "assets": "VIDEO, PDF, PPTX, DOCX, ZIP",
        "examples": [("Lesson MP4", "Video DNA"), ("Slides PPTX", "PPTX DNA")],
        "no": "Proprietary LMS packages — ZIP/PDF export if needed.",
        "media": "mixed",
    },
    {
        "name": "AI Artists",
        "role": "Generate and sell AI artworks and prompt packs.",
        "why": "Outputs and prompts are scraped, reposted, or resold without credit.",
        "help": "DNA artworks + prompt TXT/JSON; watermark commercial packs; ownership vs scrape disputes.",
        "assets": "IMAGE; TXT/JSON prompts; optional ZIP",
        "examples": [("Artwork PNG", "Image DNA"), ("Prompt JSON", "JSON DNA")],
        "no": "None for PNG/JPG + prompt LIVE types.",
        "media": "image",
    },
    {
        "name": "Fashion Designers",
        "role": "Produce lookbooks, campaigns, and line sheets for press and wholesale.",
        "why": "Campaign assets leak before launch or are stolen by counterfeit sellers.",
        "help": "DNA lookbooks/campaigns; watermark press previews; prove brand IP on stolen posts.",
        "assets": "IMAGE; VIDEO; PDF line sheets",
        "examples": [("Lookbook JPG", "Image DNA"), ("Campaign MP4", "Video DNA")],
        "no": "None for typical campaign LIVE exports.",
        "media": "image",
    },
    {
        "name": "Game Developers",
        "role": "Ship games and market them with key art, trailers, and press kits.",
        "why": "Store art and trailers are stolen for fake listings or unauthorized promos.",
        "help": "DNA key art/trailers/press kits; watermark embargo kits; ownership for art theft claims.",
        "assets": "IMAGE; VIDEO; PDF/PPTX; ZIP",
        "examples": [("Key art PNG", "Image DNA"), ("Trailer MP4", "Video DNA")],
        "no": "Game binaries (.exe/.apk) — not DNA targets; protect marketing media.",
        "media": "mixed",
    },
    {
        "name": "Influencers",
        "role": "Create and monetize social photos, reels, and brand content.",
        "why": "Reels and posts are reposted/cropped for others’ monetization.",
        "help": "DNA before post; protect brand-deal previews; monitor Shorts/repost theft; ownership evidence.",
        "assets": "VIDEO reels; IMAGE posts",
        "examples": [("Reel MP4", "Video DNA"), ("Post JPG", "Image DNA")],
        "no": "None for MP4/JPG LIVE workflow.",
        "media": "video",
    },
    {
        "name": "Advertising Agencies",
        "role": "Produce multi-channel campaigns and deliver finals to brands and publishers.",
        "why": "Pre-release creatives leak; rights disputes over who owns which asset.",
        "help": "Org vault DNA on finals; watermark pre-release shares; custody for brand legal; audit who uploaded.",
        "assets": "IMAGE, VIDEO, PDF, PPTX, ZIP",
        "examples": [("Ad MP4", "Video DNA"), ("Campaign PDF", "PDF DNA")],
        "no": "Native design tool files — export LIVE types first.",
        "media": "mixed",
    },
    {
        "name": "Creative Agencies",
        "role": "Deliver brand systems, films, and multi-format client packages.",
        "why": "Client work leaks in WIP form or is reused beyond contract scope.",
        "help": "DNA deliverables; sealed client portals; revision evolution for SOW disputes; ownership proofs.",
        "assets": "IMAGE, VIDEO, PDF, PPTX, DOCX, ZIP",
        "examples": [("Brand PDF", "PDF DNA"), ("Film MP4", "Video DNA")],
        "no": "Native creative tool files — export first.",
        "media": "mixed",
    },
    {
        "name": "Freelancers",
        "role": "Deliver contracted creative/professional finals to clients for payment.",
        "why": "Clients or third parties reuse work beyond license without credit or pay.",
        "help": "DNA invoice deliverables; watermark unpaid previews; ownership proof for license disputes.",
        "assets": "Any LIVE final (IMAGE/VIDEO/PDF/DOCX/…)",
        "examples": [("Final PNG", "Image DNA"), ("Report PDF", "PDF DNA")],
        "no": "Unsupported native finals — export to LIVE type first.",
        "media": "mixed",
    },
]

BUYERS: list[dict] = [
    {
        "name": "Marketing Teams",
        "role": "Own brand assets, campaigns, and external agency deliveries.",
        "why": "Brand kits and ads get misused externally; need proof of owned/licensed library.",
        "help": "Register brand masters; verify agency deliveries; watermark pre-launch shares; investigate misuse.",
        "assets": "IMAGE, VIDEO, PDF, PPTX, ZIP",
        "examples": [("Brand kit ZIP", "ZIP DNA"), ("Ad MP4", "Video DNA")],
        "no": "None for standard brand-kit LIVE types.",
        "media": "mixed",
    },
    {
        "name": "Advertising Agencies (Buyer role)",
        "role": "Receive creator packs and redistribute approved assets to media/publishers.",
        "why": "Need inbound authenticity and outbound leak control on client creatives.",
        "help": "DNA inbound packs; re-seal before media delivery; custody for legal sign-off; audit vendors.",
        "assets": "IMAGE, VIDEO, PDF, PPTX, ZIP",
        "examples": [("Inbound campaign ZIP", "ZIP DNA"), ("Publisher PDF", "PDF DNA")],
        "no": "Native agency tool files — LIVE export first.",
        "media": "mixed",
    },
    {
        "name": "SMEs",
        "role": "Protect core business visuals, pitch materials, and product photos.",
        "why": "Product images and pitch decks are easy to steal; limited legal budget.",
        "help": "DNA crown jewels (logo, product shots, pitch); sealed freelancer shares; ownership certificates.",
        "assets": "IMAGE; PDF; PPTX",
        "examples": [("Product JPG", "Image DNA"), ("Pitch PDF", "PDF DNA")],
        "no": "None for typical SME LIVE packs.",
        "media": "mixed",
    },
    {
        "name": "Large Enterprises",
        "role": "Govern confidential documents, media, and organization-wide digital assets.",
        "why": "Internal leaks, compliance, and executive media fraud risk at scale.",
        "help": "Full 15-layer DNA + TEP/Smart Link; custody; org audit; investigate leaks with forensic reports.",
        "assets": "All 10 LIVE types including ZIP",
        "examples": [("Board PPTX", "PPTX DNA"), ("Policy PDF", "PDF DNA"), ("Pack ZIP", "ZIP DNA")],
        "no": "XLSX not LIVE — export CSV/PDF.",
        "media": "mixed",
    },
    {
        "name": "Media Houses",
        "role": "Archive, license, and syndicate photos, video, and audio.",
        "why": "Archive assets are republished without license across the web.",
        "help": "DNA before syndication; watermark partner copies; investigate unauthorized reuse.",
        "assets": "IMAGE, VIDEO, AUDIO, PDF",
        "examples": [("Wire photo", "Image DNA"), ("Clip MP4", "Video DNA")],
        "no": "Proprietary MAM wrappers — LIVE export first.",
        "media": "mixed",
    },
    {
        "name": "Publishers",
        "role": "Publish editions, covers, and promo media under rights contracts.",
        "why": "Pirate mirrors copy covers and PDF editions.",
        "help": "DNA covers/interiors; watermark ARCs; ownership proof vs pirate sites.",
        "assets": "PDF; IMAGE; optional AUDIO/VIDEO promo",
        "examples": [("Edition PDF", "PDF DNA"), ("Cover JPG", "Image DNA")],
        "no": "InDesign/Quark native — export PDF first.",
        "media": "doc",
    },
    {
        "name": "E-commerce Brands",
        "role": "Produce and publish catalog and lifestyle product media.",
        "why": "SKU photos are scraped for fake stores and marketplace clones.",
        "help": "DNA catalog heroes; watermark wholesale lookbooks; prove ownership vs fake listings.",
        "assets": "IMAGE; VIDEO; PDF",
        "examples": [("SKU JPG", "Image DNA"), ("Lifestyle MP4", "Video DNA")],
        "no": "None for standard catalog LIVE exports.",
        "media": "image",
    },
    {
        "name": "Startups",
        "role": "Share pitch decks, demos, and product visuals with investors and press.",
        "why": "Pitch materials leak; product visuals get copied by competitors.",
        "help": "DNA pitch/demo kits; Smart Link data rooms; watermark investor previews.",
        "assets": "PPTX, PDF, VIDEO, IMAGE",
        "examples": [("Pitch PPTX", "PPTX DNA"), ("Demo MP4", "Video DNA")],
        "no": "None for standard pitch LIVE kits.",
        "media": "mixed",
    },
    {
        "name": "Educational Institutions",
        "role": "Manage curriculum, lectures, and accredited learning materials.",
        "why": "Course packs are pirated; institutional IP walks out via share links.",
        "help": "DNA official courses; watermark paid materials; custody for accreditation archives.",
        "assets": "VIDEO, PDF, PPTX, DOCX, ZIP",
        "examples": [("Lecture MP4", "Video DNA"), ("Course ZIP", "ZIP DNA")],
        "no": "Proprietary LMS formats — ZIP/PDF export if needed.",
        "media": "mixed",
    },
    {
        "name": "Government Departments",
        "role": "Issue and control official documents and public communications.",
        "why": "Controlled PDFs leak; authenticity of issued documents can be challenged.",
        "help": "DNA controlled docs; mandatory custody; watermark controlled copies; Smart Link releases.",
        "assets": "PDF, DOCX, PPTX, IMAGE, ZIP",
        "examples": [("Gazette PDF", "PDF DNA"), ("Brief DOCX", "DOCX DNA")],
        "no": "XLSX not LIVE — CSV/PDF.",
        "media": "doc",
    },
    {
        "name": "Production Houses",
        "role": "Produce masters, cuts, stills, and scripts for clients and distributors.",
        "why": "Screeners and masters leak to piracy channels.",
        "help": "Studio DNA on masters; watermark screeners; custody on distributed assets; investigate leaks.",
        "assets": "VIDEO, IMAGE, PDF, DOCX, ZIP",
        "examples": [("Cut MP4", "Video DNA"), ("Screener pack ZIP", "ZIP DNA")],
        "no": "Camera originals may need mezzanine LIVE export.",
        "media": "video",
    },
    {
        "name": "Event Companies",
        "role": "Capture and deliver event photo/video galleries to clients.",
        "why": "Galleries are commercially reposted or stolen before client delivery.",
        "help": "DNA gallery masters; watermark unpaid previews; ownership vs stolen event media.",
        "assets": "IMAGE galleries; VIDEO highlights",
        "examples": [("Gallery JPG", "Image DNA"), ("Highlight MP4", "Video DNA")],
        "no": "None for JPEG/MP4 LIVE galleries.",
        "media": "image",
    },
    {
        "name": "HR & L&D Teams",
        "role": "Publish internal training, handbooks, and policy learning packs.",
        "why": "Confidential handbooks and training leak outside the company.",
        "help": "DNA training packs; TEP/watermark confidential handbooks; custody for compliance archives.",
        "assets": "PPTX, PDF, VIDEO, DOCX",
        "examples": [("Handbook PDF", "PDF DNA"), ("Module PPTX", "PPTX DNA")],
        "no": "XLSX trackers not LIVE.",
        "media": "doc",
    },
    {
        "name": "Architects & Real Estate Firms",
        "role": "Market properties with listing photos, tours, and plan documents.",
        "why": "Stolen listing photos fuel fake property scams; plans get misused.",
        "help": "DNA listings and plans; watermark unlisted previews; prove ownership vs scam listings.",
        "assets": "IMAGE; PDF plans; VIDEO tours",
        "examples": [("Listing JPG", "Image DNA"), ("Floorplan PDF", "PDF DNA")],
        "no": "CAD natives — export PDF/PNG first.",
        "media": "mixed",
    },
    {
        "name": "NGOs",
        "role": "Create campaign media and publish sensitive program reports.",
        "why": "Campaign assets are misused; embargoed reports leak early.",
        "help": "DNA campaigns/reports; watermark embargo copies; custody for sensitive investigations.",
        "assets": "IMAGE, VIDEO, PDF, DOCX",
        "examples": [("Campaign JPG", "Image DNA"), ("Report PDF", "PDF DNA")],
        "no": "None for typical campaign LIVE packs.",
        "media": "mixed",
    },
    {
        "name": "Legal Firms",
        "role": "Collect, preserve, and exchange evidence and case documents.",
        "why": "Exhibit integrity and custody must stand up in disputes.",
        "help": "DNA exhibits; chain of custody (L13); watermark opposing-counsel copies; ownership/authenticity proofs.",
        "assets": "PDF, DOCX, IMAGE, VIDEO, ZIP",
        "examples": [("Exhibit PDF", "PDF DNA"), ("Evidence ZIP", "ZIP DNA")],
        "no": "Email .msg/.pst — export to PDF/supported type first.",
        "media": "doc",
    },
    {
        "name": "AI Companies",
        "role": "Protect demos, model cards, and proprietary sample datasets.",
        "why": "Demos and sample packs leak to partners/competitors; scrape disputes.",
        "help": "DNA demos/docs/samples; watermark partner eval packs; custody on regulated docs; ownership proofs.",
        "assets": "IMAGE, VIDEO, PDF, JSON, TXT, ZIP",
        "examples": [("Demo MP4", "Video DNA"), ("Model card PDF", "PDF DNA"), ("Sample ZIP", "ZIP DNA")],
        "no": "Model weight binaries — not DNA targets.",
        "media": "mixed",
    },
    {
        "name": "Design Studios",
        "role": "Deliver multi-format creative packages to brand clients.",
        "why": "WIP and finals leak; license scope disputes with clients.",
        "help": "DNA finals; watermark WIP shares; custody on licensed handoffs; ownership in contracts.",
        "assets": "IMAGE, PDF, PPTX, VIDEO, ZIP",
        "examples": [("Deliverable PDF", "PDF DNA"), ("Showreel MP4", "Video DNA")],
        "no": "Native design files — LIVE export first.",
        "media": "mixed",
    },
    {
        "name": "Recruiters",
        "role": "Attract talent using employer-brand videos, career decks, and campaign creatives; coordinate agencies and job marketing. (Not primarily a résumé vault.)",
        "why": "Employer-brand videos and decks get scraped or reused by other firms/agencies without permission; brand hiring campaigns lose uniqueness.",
        "help": "Register EB video/deck/images → DNA + certificate → share kits via protected link → if stolen, investigate with forensic match. Optional: DNA confidential offer PDFs — not hundreds of résumés.",
        "assets": "VIDEO, PPTX, PDF, IMAGE; optional DOCX/PDF confidential docs",
        "examples": [
            ("Career video MP4", "Video DNA"),
            ("EB deck PPTX/PDF", "PPTX/PDF DNA"),
            ("Brand still PNG/JPG", "Image DNA"),
        ],
        "no": "ATS proprietary formats unsupported. Résumé piles: technically possible as DOCX/PDF but not the primary PINIT use — protect brand kits.",
        "media": "mixed",
    },
]


def layer_how(num: int, assets: str, media: str) -> str:
    """How this layer is generated for the segment's assets (authenticated Generate DNA)."""
    base = f"On their LIVE assets ({assets}): "
    mapping = {
        1: base + "UniversalFileRouter → type engine SHA-256 (IMAGE also normalized pixel hash).",
        2: base + "Sobel edges (IMAGE) or type structure hash (PDF layout / video container / audio structure).",
        3: base + "IMAGE pHash or universal SimHash — finds crop/compress/re-encode copies of their files.",
        4: base + "Histograms (IMAGE) or type semantic/distribution fingerprint for lookalike detection.",
        5: base + "EXIF/IPTC (photos) or container/OPC metadata hash (PDF/DOCX/PPTX/audio/video).",
        6: base + "IMAGE LSB+HMAC seal, or HMAC over L1–L5 for docs/video/audio.",
        7: base + "BehavioralLayer hashes upload session (who in the team uploaded this asset).",
        8: base + "RelationshipLayer links related vault items (e.g. video ↔ still ↔ deck).",
        9: base + "OriginLayer records IP/UA/geo context of the registration upload.",
        10: base + "EvolutionLayer starts version chain (v1 → revised deliverable).",
        11: base + (
            "processLayer11 deepfake heuristics (stronger on image/video MIME)."
            if media in ("image", "video", "mixed")
            else "processLayer11 runs with owner; media heuristics strongest if they also register cover IMAGE."
        ),
        12: base + "processLayer12 stores owner watermark record (ownerUserId + dnaRecordId); use Protect Download/TEP when sharing.",
        13: base + "processLayer13 writes custody entry actor=ownerUserId for legal evidence timeline.",
        14: base + "processLayer14 commitment H(secret||fileHash||ownerUserId) for ownership proof packs.",
        15: base + "processLayer15 binds user faceEmbedding if registered; else NOT_REGISTERED still recorded.",
    }
    return mapping[num]


def add_examples_table(doc, examples: list[tuple[str, str]]):
    t = doc.add_table(rows=1 + len(examples), cols=3)
    t.style = "Table Grid"
    headers = ["Their asset", "PINIT engine", "15 layers?"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=SMALL, color=WHITE)
        shade_cell(cell, "007A8C")
    for i, (asset, engine) in enumerate(examples, start=1):
        t.rows[i].cells[0].text = asset
        t.rows[i].cells[1].text = engine
        t.rows[i].cells[2].text = "Yes — if logged in"
    add_para(doc, "", space_after=4)


def add_dna_table(doc, assets: str, media: str):
    t = doc.add_table(rows=1 + 15, cols=3)
    t.style = "Table Grid"
    headers = ["Layer", "Yes / No", "How it generates for their assets"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=SMALL, color=WHITE)
        shade_cell(cell, "1A2B4A")
    for i, (num, name, purpose) in enumerate(LAYERS):
        r = t.rows[i + 1]
        r.cells[0].text = ""
        p = r.cells[0].paragraphs[0]
        run = p.add_run(f"L{num} — {name}")
        set_run(run, bold=True, size=SMALL, color=DARK)
        p2 = r.cells[0].add_paragraph()
        run2 = p2.add_run(purpose)
        set_run(run2, size=Pt(8), color=TEAL)

        r.cells[1].text = ""
        run = r.cells[1].paragraphs[0].add_run("Yes")
        set_run(run, bold=True, size=SMALL, color=DARK)

        r.cells[2].text = ""
        run = r.cells[2].paragraphs[0].add_run(layer_how(num, assets, media))
        set_run(run, size=Pt(8), color=DARK)

    for row in t.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(0.7)
        row.cells[2].width = Inches(4.3)


def add_segment(doc, seg: dict, index: int, kind: str):
    add_para(doc, f"{index}. {seg['name']} ({kind})", bold=True, size=H3, color=NAVY, space_after=6)
    add_label_value(doc, "Responsibilities", seg["role"])
    add_label_value(doc, "Why they use PINIT", seg["why"])
    add_label_value(doc, "How PINIT helps", seg["help"])
    add_label_value(doc, "Assets on PINIT", seg["assets"])
    add_label_value(doc, "If No", seg["no"])
    add_para(doc, "Platform path: Register/Login (Pinit ID) → Upload from device → Generate DNA → Vault + Certificate → Protected share / Monitor → Investigate", size=Pt(8), color=TEAL, space_after=6)
    add_para(doc, "Their files → engine", bold=True, size=SMALL, color=NAVY, space_after=3)
    add_examples_table(doc, seg["examples"])
    add_para(doc, "15-layer DNA generation for their assets (authenticated LIVE upload)", bold=True, size=SMALL, color=NAVY, space_after=3)
    add_dna_table(doc, seg["assets"], seg["media"])
    doc.add_paragraph()


def build() -> Path:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.65)
        s.right_margin = Inches(0.65)

    add_para(
        doc,
        "PINIT-DNA — How We Secure 40 Segments",
        bold=True,
        size=H1,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        f"World presentation  ·  Creator & Buyer segments  ·  {date.today().isoformat()}",
        size=Pt(11),
        color=TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(
        doc,
        "PINIT does not run a different DNA engine per job title. Photographers and Recruiters use the same Universal DNA Engine. "
        "What changes is which assets they register and why. Every authenticated Generate DNA on a LIVE file type "
        "(IMAGE, VIDEO, AUDIO, PDF, DOCX, PPTX, TXT, CSV, JSON, ZIP) produces all 15 layers.",
        size=SMALL,
        space_after=8,
    )

    # Index
    add_para(doc, "Segment index", bold=True, size=H2, color=NAVY, space_after=4)
    idx = doc.add_table(rows=1 + len(CREATORS) + len(BUYERS), cols=3)
    idx.style = "Table Grid"
    for i, h in enumerate(["#", "Segment", "Type"]):
        cell = idx.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, bold=True, size=SMALL, color=WHITE)
        shade_cell(cell, "1A2B4A")
    row_i = 1
    for i, s in enumerate(CREATORS, start=1):
        idx.rows[row_i].cells[0].text = str(i)
        idx.rows[row_i].cells[1].text = s["name"]
        idx.rows[row_i].cells[2].text = "Creator"
        row_i += 1
    for i, s in enumerate(BUYERS, start=1):
        idx.rows[row_i].cells[0].text = str(len(CREATORS) + i)
        idx.rows[row_i].cells[1].text = s["name"]
        idx.rows[row_i].cells[2].text = "Buyer"
        row_i += 1

    doc.add_page_break()
    add_para(doc, "Part A — Creator Segments", bold=True, size=H2, color=NAVY, space_after=8)
    for i, seg in enumerate(CREATORS, start=1):
        add_segment(doc, seg, i, "Creator")
        if i % 2 == 0 and i < len(CREATORS):
            doc.add_page_break()

    doc.add_page_break()
    add_para(doc, "Part B — Buyer Segments", bold=True, size=H2, color=NAVY, space_after=8)
    for i, seg in enumerate(BUYERS, start=1):
        add_segment(doc, seg, len(CREATORS) + i, "Buyer")
        if i % 2 == 0 and i < len(BUYERS):
            doc.add_page_break()

    out = Path(__file__).resolve().parent / "PINIT-DNA_Segment_15_Layer_Code_Capability.docx"
    tmp = out.with_suffix(".docx.tmp")
    doc.save(tmp)
    try:
        if out.exists():
            out.unlink()
        tmp.replace(out)
    except PermissionError:
        raise PermissionError(f"Close Word and re-run. Temp at: {tmp}") from None
    return out


if __name__ == "__main__":
    print(f"Creators={len(CREATORS)} Buyers={len(BUYERS)} Total={len(CREATORS)+len(BUYERS)}")
    path = build()
    print(f"Wrote {path}")
