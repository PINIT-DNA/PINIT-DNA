"""Convert JUNE-2026-DEVELOPMENT-REPORT.md to Word (.docx). Documentation utility only."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT = "Times New Roman"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(18)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)


def set_run_font(run, *, bold: bool = False, italic: bool = False, size=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True, size=BODY_SIZE)
        else:
            run = p.add_run(part)
            set_run_font(run, size=BODY_SIZE)
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) or c == "" for c in cells)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = RGBColor(0, 0, 0)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            for run in p.runs:
                set_run_font(run, bold=True, size=H1_SIZE)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, bold=True, size=H2_SIZE)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            for run in p.runs:
                set_run_font(run, bold=True, size=H3_SIZE)
            i += 1
            continue

        if stripped.startswith("> "):
            add_formatted_paragraph(doc, stripped[2:], style="Intense Quote")
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            header = parse_table_row(stripped)
            i += 1
            if i < len(lines) and is_separator_row(parse_table_row(lines[i].strip())):
                i += 1
            rows: list[list[str]] = [header]
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = parse_table_row(lines[i].strip())
                if not is_separator_row(row):
                    rows.append(row)
                i += 1
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text))
                    set_run_font(run, bold=(r_idx == 0), size=Pt(10))
            doc.add_paragraph()
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, bold=True, size=BODY_SIZE)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=BODY_SIZE)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            set_run_font(run, size=BODY_SIZE)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, size=Pt(9))
            run.font.name = "Courier New"
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).resolve().parent
    md = base / "JUNE-2026-DEVELOPMENT-REPORT.md"
    out = base / "JUNE-2026-DEVELOPMENT-REPORT.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    convert(md, out)
