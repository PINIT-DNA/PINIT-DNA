#!/usr/bin/env python3
"""Generate PinIT Hub Extension — Chrome publish, usage, platforms, and development process Word doc."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "PinIT_Hub_Extension_Chrome_Publish_Usage_Development_Guide.docx"


def set_run_font(run, name="Calibri", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def hline(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B4F72")
    pBdr.append(bottom)
    pPr.append(pBdr)


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=22, bold=True, color=RGBColor(0x1B, 0x4F, 0x72))


def subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=11, color=RGBColor(0x5D, 0x6D, 0x7E))


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    for run in p.runs:
        set_run_font(run, size=11)


def numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size=11)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            set_run_font(r, size=10, bold=(i == 0))
            if i == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D6EAF8")
                shading.set(qn("w:val"), "clear")
                cell._tePr = cell._tc.get_or_add_tcPr()
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title(doc, "PinIT Hub Browser Extension")
    subtitle(doc, "Chrome / Edge Publish Guide · User Guide · Supported Apps · Development Process")
    subtitle(doc, "PinIT Hub · Publish Guardian · Document version 1.0 · July 2026")
    hline(doc)

    para(
        doc,
        "This document explains what the PinIT Hub extension is, why it was built, how users install and use it, "
        "which applications it works with, how to publish it to the Chrome Web Store (and Edge Add-ons), "
        "and how the engineering team developed it—including the real product limitations (what it does and does not connect to).",
    )

    # ── 1 ──────────────────────────────────────────────────────────────────
    heading(doc, "1. Executive summary", 1)
    para(
        doc,
        "PinIT Hub is a Manifest V3 browser extension for Google Chrome and Microsoft Edge. "
        "Its product mode is Publish Guardian: when a signed-in user publishes or exports media on supported websites, "
        "the extension can capture the original file, send it to PinIT Hub’s backend, store it in the Vault, "
        "generate DNA fingerprints, issue a certificate, enroll monitoring, and show status in the Hub UI "
        "(Protected Posts, Assets, Monitoring, Investigation).",
    )
    para(doc, "Product promise (approved wording):", bold=True)
    para(
        doc,
        "Install PinIT once — when you post on supported sites in Chrome, your content is automatically protected "
        "in Vault, DNA-registered, and monitored for public copies and tampering, with alerts in PinIT Hub.",
    )
    para(doc, "Important clarifications:", bold=True)
    bullet(doc, "The extension connects to the user’s PinIT Hub account (JWT via one-time auth code).")
    bullet(
        doc,
        "It does NOT OAuth-connect Instagram, LinkedIn, YouTube, or Google accounts. "
        "It does not receive platform passwords.",
    )
    bullet(
        doc,
        "Protection happens because the extension runs inside the user’s browser and can see files they choose "
        "to upload / images they protect—not because PinIT logs into those platforms as the user.",
    )
    bullet(
        doc,
        "Right-click Protect is a fallback. The intended UX is automatic protect-on-publish / protect-on-export; "
        "reliability varies by site (stronger on file-picker uploads, weaker on sites that hijack right-click, e.g. YouTube Shorts player).",
    )

    # ── 2 ──────────────────────────────────────────────────────────────────
    heading(doc, "2. Why we built the extension (business & technical uses)", 1)
    heading(doc, "2.1 Business uses", 2)
    bullet(doc, "Protect creator and enterprise media at the moment of publish—not only after a leak is found.")
    bullet(doc, "Give users one Hub dashboard for Vault ownership, certificates, monitoring, and investigations.")
    bullet(doc, "Reduce friction: users keep posting on Instagram / LinkedIn / YouTube Studio / Canva as usual.")
    bullet(doc, "Support Brand / IP / photographer workflows: portfolio sites, export tools, CMS uploads.")
    bullet(doc, "Differentiate PinIT Hub as “protection in the browser” plus “forensics in the Hub.”")

    heading(doc, "2.2 Technical uses", 2)
    bullet(doc, "Capture original bytes before social platforms recompress or strip metadata.")
    bullet(doc, "Reuse existing Hub engines (DNA, Vault, Certificates, Monitoring, Investigation)—no parallel stack.")
    bullet(doc, "Work cross-site without each platform’s partner API for basic protect-on-upload.")
    bullet(doc, "Provide Verify (match against Vault) and Protect (enroll new asset) from any page.")
    bullet(doc, "Queue / retry when the network or Render backend is cold-starting.")

    heading(doc, "2.3 What the extension is not", 2)
    bullet(doc, "Not a replacement for YouTube/Instagram official APIs or channel management tools.")
    bullet(doc, "Not continuous reading of private DMs, WhatsApp, Discord private chats.")
    bullet(doc, "Not a guarantee of 100% leak detection on the entire internet.")
    bullet(doc, "Not silent account takeover of social platforms.")

    # ── 3 ──────────────────────────────────────────────────────────────────
    heading(doc, "3. How a user installs and uses the extension", 1)
    heading(doc, "3.1 Today (developer / sideload — current team testing path)", 2)
    numbered(doc, "Open chrome://extensions (Chrome) or edge://extensions (Edge).")
    numbered(doc, "Turn on Developer mode.")
    numbered(doc, "Click Load unpacked → select the project’s extension/ folder.")
    numbered(doc, "Open Extension options → set API and Hub URLs → Save.")
    numbered(doc, "Click the PinIT Hub icon → Sign in to PinIT.")
    numbered(doc, "Authorize on Hub (www.pinithub.com/extension/auth) → copy one-time code → Connect in the popup.")
    numbered(doc, "Confirm status shows CONNECTED.")

    para(doc, "Production defaults (company Hub):", bold=True)
    add_table(
        doc,
        [
            ["Setting", "Value"],
            ["API Base URL", "https://pinit-dna-uf5y.onrender.com/api/v1"],
            ["Hub Base URL", "https://www.pinithub.com"],
            ["Auth page", "https://www.pinithub.com/extension/auth"],
            ["Results UI", "Protected Posts, Assets, Vault, Monitoring"],
        ],
    )

    heading(doc, "3.2 How users will install after Chrome Web Store publish", 2)
    numbered(doc, "Open the Chrome Web Store listing for “PinIT Hub”.")
    numbered(doc, "Click Add to Chrome (or Get on Edge Add-ons).")
    numbered(doc, "Confirm permissions.")
    numbered(doc, "Pin the extension → Sign in with PinIT Hub account (same auth-code flow).")
    numbered(doc, "Enable Publish Guardian and desired platforms in Options.")
    numbered(doc, "Post / export normally on supported sites; check Hub for Vault / Protected Posts / Assets.")

    heading(doc, "3.3 Everyday usage modes", 2)
    add_table(
        doc,
        [
            ["Mode", "How the user triggers it", "What happens"],
            [
                "Publish Guardian (auto)",
                "Choose a file to upload/export on a supported site (e.g. YouTube Studio, Canva export)",
                "Extension captures bytes → publish-protect API → Vault + DNA + Cert + Monitor + Asset/Post records",
            ],
            [
                "Protect (manual)",
                "Right-click image (or Shift+Right-click on stubborn players) → Protect with PinIT",
                "Same protect pipeline using image/video URL or page thumbnail",
            ],
            [
                "Verify",
                "Right-click image → Verify with PinIT",
                "Compares media against Vault identity / DNA match",
            ],
            [
                "Hub Protect file",
                "Upload in www.pinithub.com → Protect file",
                "Same backend protection without the extension",
            ],
            [
                "Monitor & investigate",
                "Open Protected Posts / Assets / Monitoring / Investigation in Hub",
                "Lifecycle timeline, discoveries, tamper signals, certificates",
            ],
        ],
    )

    heading(doc, "3.4 End-to-end protect flow (happy path)", 2)
    para(
        doc,
        "User signed into PinIT Hub via extension → selects media on a supported site → extension captures "
        "original media + metadata (platform, page URL, caption when available) → POST /api/v1/extension/publish-protect "
        "→ backend generates DNA, stores Vault, issues certificate (if entitled), creates/links Asset and Protected Post, "
        "enrolls monitoring → extension badge shows success → user opens Hub to see timeline and status.",
    )

    heading(doc, "3.5 Continuous tracking (“where is my content going?”)", 2)
    para(
        doc,
        "After protect, tracking is not “live login to YouTube/Instagram.” It is Hub monitoring enrolled on the "
        "protected asset and watch URLs (post URL / profile URL when known), plus DNA matching when the crawler "
        "or monitors find candidate copies. Discoveries and risk appear on Protected Posts / Assets / Monitoring. "
        "Users can open Unified Investigation for deeper forensics.",
    )

    # ── 4 ──────────────────────────────────────────────────────────────────
    heading(doc, "4. Applications / platforms the extension works with", 1)
    para(
        doc,
        "Platforms are registered in extension/shared/platforms.js and injected via content scripts in manifest.json. "
        "Each site adapter listens for file-input / export patterns where possible. Users can toggle platforms in Options.",
    )

    heading(doc, "4.1 Social", 2)
    bullet(doc, "Instagram (Web)")
    bullet(doc, "Facebook (Web)")
    bullet(doc, "X / Twitter (Web)")
    bullet(doc, "Pinterest (Web)")
    bullet(doc, "Telegram Web (public content only)")
    bullet(doc, "YouTube / YouTube Studio")
    bullet(doc, "TikTok Web")

    heading(doc, "4.2 Creators & publishing", 2)
    bullet(doc, "Threads, Reddit, Tumblr, Medium, Substack, Patreon")
    bullet(doc, "Vimeo, Twitch")
    bullet(doc, "Behance, Dribbble, ArtStation, DeviantArt")

    heading(doc, "4.3 Photography / portfolio", 2)
    bullet(doc, "Flickr, 500px, SmugMug, Pixieset, Zenfolio")
    bullet(doc, "Adobe Portfolio, Squarespace, Wix")

    heading(doc, "4.4 Business / design / commerce", 2)
    bullet(doc, "LinkedIn")
    bullet(doc, "GitHub")
    bullet(doc, "Canva, Figma, Adobe Express, Photoshop Web, Illustrator Web (upload + export protect paths)")
    bullet(doc, "Shopify Admin, WordPress Admin")

    heading(doc, "4.5 How protection works per application type", 2)
    add_table(
        doc,
        [
            ["Application type", "Typical trigger", "Notes"],
            [
                "Social / creators with file upload",
                "User selects image/video in upload UI",
                "File-input adapter captures bytes in parallel; platform upload continues normally",
            ],
            [
                "Design tools (Canva / Figma / Adobe Web)",
                "Export / download / publish actions",
                "Export-protect script offers parallel protect without blocking export",
            ],
            [
                "Any website with an image",
                "Right-click → Protect / Verify",
                "Uses image URL; works even if site has no dedicated adapter",
            ],
            [
                "YouTube Shorts watch page",
                "Weak for normal right-click (player menu)",
                "Prefer Studio upload, Shift+Right-click page Protect, or Hub Protect file with original video",
            ],
            [
                "Already-live posts without capture",
                "Manual Protect or Hub upload",
                "No platform account sync—extension cannot list “all my channel videos” without future OAuth",
            ],
        ],
    )

    heading(doc, "4.6 Out of scope (by design)", 2)
    bullet(doc, "Private WhatsApp / Discord / Messenger DMs")
    bullet(doc, "Native mobile apps (extension is browser-only)")
    bullet(doc, "Platforms without a web surface the extension can see")

    # ── 5 ──────────────────────────────────────────────────────────────────
    heading(doc, "5. Publishing the extension to Chrome (and Edge)", 1)
    heading(doc, "5.1 Package the extension", 2)
    numbered(doc, "Ensure manifest.json name/version/description/icons are final (e.g. PinIT Hub, version 1.2.x).")
    numbered(doc, "Set production defaults in shared/config.js (API = Render Hub backend, Hub = www.pinithub.com).")
    numbered(doc, "Remove debug-only code; confirm host_permissions match production domains.")
    numbered(doc, "Zip the contents of the extension/ folder (not the parent repo). Include manifest, background, content, popup, options, icons, shared.")
    numbered(doc, "Sideload-test the zip/unpacked build on a clean Chrome profile end-to-end.")

    heading(doc, "5.2 Chrome Web Store publish steps", 2)
    numbered(doc, "Create a Google Chrome Web Store Developer account and pay the one-time developer fee.")
    numbered(doc, "Open Chrome Web Store Developer Dashboard → New item → upload the .zip.")
    numbered(doc, "Complete store listing: title, short/long description, screenshots (1280×800 or required sizes), category, language.")
    numbered(doc, "Privacy practices: explain what data is collected (media chosen for protect, Hub auth tokens stored locally, page URLs for monitoring enrollment).")
    numbered(doc, "Justify host permissions / <all_urls> if used (needed for Protect/Verify on arbitrary sites and adapters).")
    numbered(doc, "Single purpose statement: digital asset protection / publish-time vaulting and verification for PinIT Hub.")
    numbered(doc, "Submit for review. Address policy questions (remote code, broad permissions, user data).")
    numbered(doc, "After approval, users install via Add to Chrome. Ship updates by uploading a new zip with a bumped version.")

    heading(doc, "5.3 Microsoft Edge Add-ons", 2)
    para(
        doc,
        "Edge can install Chrome Web Store extensions, or you publish separately to Microsoft Edge Add-ons "
        "(Partner Center). Use the same MV3 package; test auth with Edge’s extension:// origin (CORS already allows extension:// on the Hub backend).",
    )

    heading(doc, "5.4 Store listing suggested copy (short)", 2)
    para(
        doc,
        "PinIT Hub protects your images and videos when you publish online. Connect your PinIT Hub account once, "
        "then publish on supported sites. PinIT can vault originals, register DNA identity, issue certificates, "
        "and monitor for public copies—then show activity in PinIT Hub.",
    )

    heading(doc, "5.5 Compliance checklist before submit", 2)
    bullet(doc, "Privacy policy URL on pinithub.com")
    bullet(doc, "Clear disclosure: media is uploaded to PinIT servers when user protects/publishes with extension enabled")
    bullet(doc, "No remote-hosted executable extension code")
    bullet(doc, "Auth tokens only in extension storage; use HTTPS APIs")
    bullet(doc, "Support contact / Hub help link")

    # ── 6 ──────────────────────────────────────────────────────────────────
    heading(doc, "6. Architecture (how the pieces connect)", 1)
    add_table(
        doc,
        [
            ["Layer", "Location", "Role"],
            ["Extension UI", "extension/popup, options", "Sign-in, status, platform toggles, deep links to Hub"],
            ["Service worker", "extension/background", "Context menus, queue flush, API calls, badges"],
            ["Content scripts / adapters", "extension/content", "Per-site file capture & export hooks"],
            ["Hub frontend", "client/ (Vercel / pinithub.com)", "Authorize extension, Protected Posts, Assets, Vault UI"],
            ["Backend API", "src/ on Render (pinit-dna-uf5y)", "publish-protect, register-post, extension auth, posts, assets"],
            ["Data", "Supabase PostgreSQL + Storage", "Users, Vault, DNA, certificates, protected posts, assets, auth codes"],
            ["Monitoring / AI", "existing Hub crawler + python-ai", "Discovery, matching, investigation support"],
        ],
    )

    heading(doc, "6.1 Extension authentication (no platform OAuth)", 2)
    numbered(doc, "Popup opens Hub /extension/auth?ext_id=<chrome.runtime.id>.")
    numbered(doc, "Logged-in Hub user clicks Authorize → backend stores one-time code bound to that extension ID (≈5 minutes, single use).")
    numbered(doc, "User pastes code in popup → POST /auth/extension/token → Hub JWTs stored in chrome.storage.")
    numbered(doc, "Later API calls use Bearer token. CORS allows chrome-extension:// and extension:// origins.")

    heading(doc, "6.2 Core API endpoints used by the extension", 2)
    bullet(doc, "POST /api/v1/auth/extension/issue-code (Hub session)")
    bullet(doc, "POST /api/v1/auth/extension/token (exchange code)")
    bullet(doc, "POST /api/v1/extension/publish-protect (multipart media + metadata)")
    bullet(doc, "POST /api/v1/extension/register-post")
    bullet(doc, "GET /api/v1/extension/sync")
    bullet(doc, "POST /api/v1/vault/verify-identity (Verify mode)")

    # ── 7 ──────────────────────────────────────────────────────────────────
    heading(doc, "7. How we developed it (process)", 1)
    heading(doc, "7.1 Development approach", 2)
    para(
        doc,
        "The extension was built as an additive module on PinIT Hub—not a rewrite of Vault, DNA, Monitoring, or Investigation. "
        "Work landed primarily on the ashwitha branch of the company repository (PINIT-DNA/PINIT-DNA), with Render backend "
        "deployed from that branch and Hub frontend on www.pinithub.com.",
    )

    heading(doc, "7.2 Phased delivery", 2)
    add_table(
        doc,
        [
            ["Phase", "Focus", "Outcome"],
            ["Phase 1 — Verify", "Right-click verify / investigate entry points", "Users can check if media matches Vault identity"],
            [
                "Phase 2 — Publish Guardian",
                "Protect-on-publish, Protected Posts UI, extension auth, queue/retry",
                "Protect creates Vault+DNA+Cert+Monitor records visible in Hub",
            ],
            [
                "Asset-first layer",
                "Asset model linked to Protected Posts; Assets dashboard",
                "Protect also creates/links Assets; lifecycle timeline",
            ],
            [
                "Platform pack",
                "Many site adapters + photographer/export tools + Options toggles",
                "Coverage across social, creator, photo, business tools",
            ],
            [
                "Reliability / Edge",
                "CORS for extension://, production API host, auth-code table ensure, clearer errors",
                "Edge Connect flow works against live Hub",
            ],
        ],
    )

    heading(doc, "7.3 Engineering process (how the team worked)", 2)
    numbered(doc, "Specify product promise and non-goals (Publish Guardian spec).")
    numbered(doc, "Add backend routes/services (extension auth, publish-protect) reusing Vault/DNA/monitor services.")
    numbered(doc, "Add Prisma models/migrations (Protected Posts, ExtensionAuthCode, Assets) + ensure scripts for production boot.")
    numbered(doc, "Build MV3 extension package (popup, options, service worker, adapters, queue).")
    numbered(doc, "Build Hub pages (/extension/auth, /protected-posts, /assets).")
    numbered(doc, "Local test with Load unpacked + localhost or production API.")
    numbered(doc, "Push to company git remote branch ashwitha → Render auto-deploy backend → verify live API.")
    numbered(doc, "Fix deploy blockers (TypeScript on Render/Vercel, CORS, wrong API host, missing DB table).")
    numbered(doc, "Validate full path on Edge: Connect → Protect → Vault / Protected Posts / Assets.")
    numbered(doc, "Document gaps (no YouTube OAuth; Shorts right-click; auto vs manual) for honest product communication.")

    heading(doc, "7.4 Key repository paths", 2)
    bullet(doc, "extension/ — MV3 package (ship this folder to the Store)")
    bullet(doc, "src/services/publish-guardian/ — protect + extension auth services")
    bullet(doc, "src/api/routes/publish-guardian.routes.ts — HTTP routes")
    bullet(doc, "client/src/pages/publish-guardian/ — Hub UI for posts + extension auth")
    bullet(doc, "client/src/pages/assets/ — Assets UI")
    bullet(doc, "docs/PINIT_Publish_Guardian_Spec.md — architecture spec")

    heading(doc, "7.5 Testing checklist used during development", 2)
    bullet(doc, "Extension Connect (auth code) on Edge and Chrome")
    bullet(doc, "Options API URL points at live Render host")
    bullet(doc, "Protect image → appears in Vault + Protected Posts + Assets")
    bullet(doc, "Certificate + monitoring status visible")
    bullet(doc, "Queue retry when backend cold-starts")
    bullet(doc, "Platform toggle OFF skips capture")
    bullet(doc, "YouTube: Studio upload path preferred over Shorts player right-click")

    # ── 8 ──────────────────────────────────────────────────────────────────
    heading(doc, "8. Current limitations & next roadmap", 1)
    heading(doc, "8.1 Known limitations", 2)
    bullet(doc, "Not published to Chrome Web Store yet (sideload / unpacked for testing).")
    bullet(doc, "No official OAuth link to YouTube/Instagram/LinkedIn accounts—no auto-import of entire channel history.")
    bullet(doc, "Auto-protect reliability differs by site; right-click remains necessary as fallback.")
    bullet(doc, "YouTube Shorts player replaces the browser context menu unless Shift+Right-click / page menu / Studio upload.")
    bullet(doc, "Large video protect may hit size/timeouts depending on backend limits.")
    bullet(doc, "Private messaging apps are out of scope.")

    heading(doc, "8.2 Recommended next steps", 2)
    bullet(doc, "Chrome Web Store + Edge Add-ons listing and screenshots.")
    bullet(doc, "Strengthen auto-capture on top platforms (LinkedIn, Instagram, YouTube Studio).")
    bullet(doc, "Optional “Connect YouTube” OAuth for channel-level enrollment (future).")
    bullet(doc, "Clearer labels (profile photo vs feed post vs Short).")
    bullet(doc, "In-Hub image/video preview on every Protected Post (thumbnail from Vault).")
    bullet(doc, "User education in popup: auto vs manual, what “CONNECTED” means.")

    # ── 9 ──────────────────────────────────────────────────────────────────
    heading(doc, "9. Quick reference URLs", 1)
    add_table(
        doc,
        [
            ["Purpose", "URL"],
            ["Hub", "https://www.pinithub.com"],
            ["Extension authorize", "https://www.pinithub.com/extension/auth"],
            ["Protected Posts", "https://www.pinithub.com/protected-posts"],
            ["Assets", "https://www.pinithub.com/assets"],
            ["Vault", "https://www.pinithub.com/vault"],
            ["API", "https://pinit-dna-uf5y.onrender.com/api/v1"],
            ["Local extension load", "chrome://extensions or edge://extensions → Load unpacked → extension/"],
        ],
    )

    heading(doc, "10. Closing statement", 1)
    para(
        doc,
        "The PinIT Hub extension exists so protection travels with the user’s publishing workflow in the browser. "
        "Users connect once to PinIT Hub—not to each social network. Media they protect (automatically on supported "
        "uploads/exports, or manually via Protect) is vaulted, DNA-registered, certificated, and monitored, with "
        "full activity visible in PinIT Hub. Publishing to the Chrome Web Store is the distribution step that turns "
        "today’s developer sideload into a one-click Get experience for customers; the protect/track engines behind "
        "that button are already implemented on the Hub stack and validated on live production paths.",
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
