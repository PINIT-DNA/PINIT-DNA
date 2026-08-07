#!/usr/bin/env python3
"""Generate PinIT Hub Video Test & Demo Recording Guide as Word .docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "PinIT_Hub_Video_Test_and_Demo_Recording_Guide.docx"
# Fallback if the main file is locked (open in Word / OneDrive)
ALT_OUT_PATH = ROOT / "docs" / "PinIT_Hub_Video_Test_and_Demo_Recording_Guide_v11.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY = RGBColor(0x1B, 0x4F, 0x72)
DARK = RGBColor(0x2C, 0x3E, 0x50)


def set_run_font(run, name="Calibri", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color or BLACK


def add_hline(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B4F72")
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "1B4F72")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else DARK
    return h


def add_p(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=size)


def add_numbered(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=size)


def add_table(doc, rows, col_widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            is_header = header and i == 0
            set_run_font(run, size=10, bold=is_header, color=RGBColor(0xFF, 0xFF, 0xFF) if is_header else BLACK)
            set_cell_border(cell)
            if is_header:
                shade_cell(cell, "1B4F72")
            elif i % 2 == 0:
                shade_cell(cell, "F4F8FB")
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True, color=NAVY)
    add_p(doc, body, size=11)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PinIT Hub")
    set_run_font(run, size=28, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Video Test Recording & Demo Walkthrough Guide")
    set_run_font(run, size=16, bold=True, color=DARK)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "For end-stage testing · Record every test · Proud product demo quality\n"
        "Companion to: PINIT-DNA User Test Flow (Tests 1–24)\n"
        "Live app: https://dna-pinit-web.vercel.app"
    )
    set_run_font(run, size=10, color=DARK)
    add_hline(doc)

    # ── 1. Your assignment ──
    add_heading(doc, "1. Your Assignment (What You Must Deliver)", 1)
    add_p(
        doc,
        "You are doing the final walkthrough of PinIT Hub. Your job is not only to click through "
        "the app — you must record a clear video for every test so the team can proudly show that "
        "PinIT Hub solves real ownership, vault, sharing, tracking, and investigation problems.",
    )
    add_p(doc, "Deliverables", bold=True)
    add_bullets(
        doc,
        [
            "One recorded video for every test in the User Test Flow (Test 1 through Test 24, plus End-to-End if asked).",
            "Business Account videos for Part F (Tests 25–33): Business dashboard, setup wizard, team, org profile, protect stack in Business shell.",
            "One Landing / Login demo video that introduces PinIT Hub to a new viewer.",
            "One short “proud solutions” demo video (optional but recommended) showing the full story: protect → vault → share → track → investigate (add Business shell if using a Business account).",
            "A simple pass/fail note for each test (Pass / Fail / Blocked + short comment).",
        ],
    )
    add_callout(
        doc,
        "Reference document you must follow for steps:",
        "PINIT-DNA-USER-TEST-FLOW.pdf / docs/PINIT-DNA-USER-TEST-FLOW.docx — use that document for exact click steps and expected results. This guide tells you HOW to record and present those tests as proud demos.",
    )

    # ── 2. Goal of every video ──
    add_heading(doc, "2. Goal of Every Video (Proud Demo Style)", 1)
    add_p(
        doc,
        "Each video should answer three questions for anyone watching (founder, client, investor, teammate):",
    )
    add_numbered(
        doc,
        [
            "What problem does the user face?",
            "What does PinIT Hub do to solve it?",
            "What proof do we see on screen that it worked?",
        ],
    )
    add_p(doc, "Simple speaking pattern (use this every time)", bold=True)
    add_bullets(
        doc,
        [
            "Start (5–10 sec): “Problem: …”",
            "Middle: Do the steps slowly. Say what you are clicking.",
            "End (10–15 sec): “Result: …” and point to the success on screen (ID, status, map, certificate, etc.).",
        ],
    )
    add_callout(
        doc,
        "Tone:",
        "Calm, clear, confident. Speak like a product walkthrough, not a rushed QA click-through. Pause 1 second after each important result so it is readable on video.",
    )

    # ── 3. Before you start ──
    add_heading(doc, "3. Before You Start — Setup Checklist", 1)
    add_table(
        doc,
        [
            ["Item", "What to prepare"],
            ["URL", "https://dna-pinit-web.vercel.app"],
            ["Browser", "Chrome or Edge (latest). Use a clean profile if possible."],
            ["Account", "Test credentials from the PinIT team (User ID + biometric if required)."],
            ["Location", "Allow GPS when browser asks (needed for share viewing tests)."],
            ["Camera / Mic", "Needed for login biometric tests and for your voice in demos."],
            ["Test files", "Prepare 2–3 files: 1 PDF with fake sensitive text, 1 JPG/PNG image, 1 DOCX."],
            ["Second window", "Incognito window ready for recipient share tests."],
            ["Server wake-up", "First load may take 30–60 seconds. Refresh once if empty."],
            ["Hide secrets", "Do not show real passwords, private keys, or personal Aadhaar/PAN of real people."],
        ],
        col_widths=[1.6, 5.2],
    )

    add_heading(doc, "3.1 How to record on Windows", 2)
    add_p(doc, "Recommended options (pick one and stick to it):", bold=True)
    add_bullets(
        doc,
        [
            "Windows Xbox Game Bar (easy): Press Win + G → Capture → Start recording. Or Win + Alt + R to start/stop.",
            "OBS Studio (best quality): Free, clear full-screen browser capture + microphone.",
            "Phone as backup: Only if screen recording fails — less preferred for desktop UI demos.",
        ],
    )
    add_p(doc, "Recording settings (proud quality)", bold=True)
    add_table(
        doc,
        [
            ["Setting", "Recommended"],
            ["Resolution", "1920×1080 (Full HD) if possible"],
            ["Browser zoom", "100% (or 110% if text is hard to read)"],
            ["Microphone", "On — speak clearly"],
            ["Mouse", "Move slowly; click once; avoid fast scrolling"],
            ["Notifications", "Turn off WhatsApp / email popups during recording"],
            ["Desktop", "Close unrelated tabs; only PinIT Hub visible"],
            ["Length", "Usually 1–4 minutes per test video; landing demo 2–3 minutes"],
        ],
        col_widths=[1.8, 5.0],
    )

    add_heading(doc, "3.2 Video file naming (mandatory)", 2)
    add_p(doc, "Save every file with this name pattern so the team can find videos fast:")
    add_p(doc, "PINIT_TestXX_ShortName_YYYY-MM-DD.mp4", bold=True, size=12)
    add_p(doc, "Examples:")
    add_bullets(
        doc,
        [
            "PINIT_LandingDemo_2026-08-03.mp4",
            "PINIT_Test01_SecureLogin_2026-08-03.mp4",
            "PINIT_Test03_GenerateDNA_2026-08-03.mp4",
            "PINIT_Test10_SecureShare_2026-08-03.mp4",
            "PINIT_ProudSolutions_Demo_2026-08-03.mp4",
        ],
    )
    add_p(
        doc,
        "Suggested folder: PinIT-Hub-Test-Videos / Day-1 / Day-2. Upload to the shared drive or chat the team provides.",
    )

    # ── 4. Landing / Hub demo ──
    add_heading(doc, "4. Landing Page & PinIT Hub Demo Video (Must Record)", 1)
    add_p(
        doc,
        "This is the first video a new person should watch. It introduces PinIT Hub before deep test cases. "
        "The public entry is the Login / Gateway page at /login. After login, show the Hub dashboard.",
    )
    add_p(doc, "Target length: 2 to 3 minutes", bold=True)

    add_heading(doc, "4.1 What this demo must prove", 2)
    add_bullets(
        doc,
        [
            "PinIT Hub is a secure digital vault + forensic platform (not a normal file drive).",
            "Users get identity (DNA), encrypted storage (Vault), tracked sharing, and investigation.",
            "The product looks professional and ready for real use.",
        ],
    )

    add_heading(doc, "4.2 Exact recording script (say this + do this)", 2)
    add_table(
        doc,
        [
            ["Time", "Say (voice)", "Do on screen"],
            [
                "0:00–0:15",
                "“Welcome to PinIT Hub — Secure, Connect, Control. This is our digital vault and forensic platform.”",
                "Open https://dna-pinit-web.vercel.app/login. Show brand / login screen full view.",
            ],
            [
                "0:15–0:35",
                "“Users sign in with their PinIT identity. Access is protected before any vault data is shown.”",
                "Point to login fields / biometric options. Start login (do not expose secrets).",
            ],
            [
                "0:35–1:05",
                "“After login, this is the Hub dashboard — DNA records, vault files, and forensic actions in one place.”",
                "Land on Dashboard. Slowly pan across stat cards and quick actions.",
            ],
            [
                "1:05–1:40",
                "“From here, an owner can generate DNA, open the encrypted vault, share with tracking, or start investigation.”",
                "Hover/click sidebar items briefly: Generate DNA → Vault → Access Intelligence → Unified Investigation → Certificates. Do not go deep yet.",
            ],
            [
                "1:40–2:20",
                "“PinIT Hub solves the full lifecycle: prove ownership, store securely, share with accountability, and investigate leaks.”",
                "Return to Dashboard. Hold on the overview for a clean ending frame.",
            ],
            [
                "2:20–2:40",
                "“Next videos show each module working end to end. Thank you.”",
                "Stop recording on Dashboard with brand visible.",
            ],
        ],
        col_widths=[1.1, 2.9, 2.8],
    )

    add_heading(doc, "4.3 Landing demo checklist", 2)
    add_bullets(
        doc,
        [
            "[ ] Brand / PinIT Hub visible at the start",
            "[ ] Login path shown without exposing credentials",
            "[ ] Dashboard overview clear",
            "[ ] Sidebar modules briefly introduced",
            "[ ] Voice explains the problem–solution story",
            "[ ] No console errors, blank screens, or popups",
            "[ ] File named PINIT_LandingDemo_YYYY-MM-DD.mp4",
        ],
    )

    # ── 5. Test Case 1 detail ──
    add_heading(doc, "5. Test Case 1 — Secure User Login (How to Test + How to Record)", 1)
    add_p(
        doc,
        "This is your first formal test video. Follow these steps exactly. "
        "Module: Authentication (PINIT HOID).",
    )

    add_heading(doc, "5.1 Problem we are proving", 2)
    add_p(
        doc,
        "Problem: Sensitive vaults must not open to strangers. Only a registered PinIT identity should enter the Hub.",
    )
    add_p(
        doc,
        "Solution: Secure login with PinIT User ID and verification (Face / Voice / Device as configured).",
    )

    add_heading(doc, "5.2 Prepare before recording", 2)
    add_numbered(
        doc,
        [
            "Open Chrome/Edge. Close extra tabs.",
            "Have test User ID ready (from team). Do not paste passwords into chat or leave them on screen notes.",
            "If biometric login is enabled, allow camera and sit in good light.",
            "Start screen recording (Win + Alt + R or OBS).",
            "Speak into mic: introduce Test 1.",
        ],
    )

    add_heading(doc, "5.3 Test steps (do these while recording)", 2)
    add_table(
        doc,
        [
            ["Step", "Action", "What to say"],
            ["1", "Open https://dna-pinit-web.vercel.app", "“Opening the live PinIT Hub application.”"],
            ["2", "Go to Login (/login) if not already there", "“This is the secure login gateway.”"],
            ["3", "Enter your PINIT User ID", "“Entering a registered PinIT User ID.”"],
            ["4", "Complete verification (Face / Voice / Device)", "“Completing identity verification.”"],
            ["5", "Wait for redirect to Dashboard", "“Login succeeded — redirecting to the Hub dashboard.”"],
            ["6", "Show User ID in sidebar", "“Here is my PinIT ID in the sidebar — proof we are authenticated.”"],
            ["7", "Confirm Dashboard loaded with no errors", "“Dashboard loaded. Test 1 passed.”"],
        ],
        col_widths=[0.7, 2.8, 3.3],
    )

    add_heading(doc, "5.4 Expected result (must show on camera)", 2)
    add_bullets(
        doc,
        [
            "[ ] Login succeeds only for registered users",
            "[ ] User ID appears in sidebar (example format: PINIT-XXXXXXXX)",
            "[ ] Dashboard loads without errors or blank page",
        ],
    )

    add_heading(doc, "5.5 What to do if something fails", 2)
    add_table(
        doc,
        [
            ["Issue", "What to try", "How to mark"],
            ["Page empty / spinning long", "Wait 60s for server wake-up, then refresh once", "Retry; note wake-up delay"],
            ["Login rejected", "Confirm credentials with team; retry once on camera", "Fail if still rejected"],
            ["Biometric camera blocked", "Allow camera in browser settings; retry", "Blocked if hardware unavailable"],
            ["Dashboard error toast", "Capture error text clearly on video", "Fail + paste error in notes"],
        ],
        col_widths=[2.0, 2.8, 2.0],
    )

    add_heading(doc, "5.6 After Test 1 video", 2)
    add_bullets(
        doc,
        [
            "Stop recording.",
            "Save as: PINIT_Test01_SecureLogin_YYYY-MM-DD.mp4",
            "Write result: Pass / Fail / Blocked + 1-line comment.",
            "Stay logged in for Test 2 (Dashboard).",
        ],
    )

    # ── 6. Pattern for all other tests ──
    add_heading(doc, "6. How to Record Every Other Test (Same Pattern)", 1)
    add_p(
        doc,
        "For Tests 2–24, open the User Test Flow document and follow its steps. "
        "Use this recording template every time so videos stay consistent and proud.",
    )
    add_p(doc, "30-second opening (always)", bold=True)
    add_numbered(
        doc,
        [
            "Say: “This is Test X — [Module name].”",
            "Say the user problem in one sentence.",
            "Say what PinIT Hub will prove in this video.",
        ],
    )
    add_p(doc, "During steps", bold=True)
    add_bullets(
        doc,
        [
            "Follow the numbered steps from the User Test Flow.",
            "Narrate important clicks: “Now I upload…”, “Now I create share link…”, “Now I open Access Intelligence…”",
            "When a success ID appears (DNA ID, Vault ID, Certificate ID, share URL), pause and read it aloud.",
        ],
    )
    add_p(doc, "Ending (always)", bold=True)
    add_bullets(
        doc,
        [
            "Tick expected results verbally: “Expected result one — visible. Expected result two — visible.”",
            "Say: “Test X — Pass” (or Fail/Blocked).",
            "End on a clean success screen.",
        ],
    )

    add_heading(doc, "6.1 Problem → Solution lines for key tests (say these)", 2)
    add_table(
        doc,
        [
            ["Test", "Problem (say this)", "Solution proof (show this)"],
            ["2 Dashboard", "Owners need one place to see vault health and actions.", "Stats, charts, quick actions work."],
            ["3 Generate DNA", "Filenames are weak ownership proof.", "DNA fingerprint completes; DNA Record ID shown."],
            ["4 Duplicate", "Same file re-uploaded creates confusion.", "Duplicate warning / forensic log appears."],
            ["5–7 Vault", "Originals need encrypted storage and safe retrieve.", "Vault ID, AES encryption, retrieve works."],
            ["9 Timeline", "No chain of custody for file events.", "Events show time, device, location."],
            ["10–12 Share", "Normal links give zero accountability.", "Tracked link, GPS, watermark, viewer activity."],
            ["14 Protected DL", "Clean downloads leave no trail.", "TEP / protected package + timeline event."],
            ["15 Access Intel", "Owner cannot see who viewed where.", "Map, viewers, revoke works."],
            ["16 Intel Report", "Need one forensic summary.", "Identity, integrity, distribution, risk sections."],
            ["18 Investigation", "Leaked copy — who owns the original?", "Owner recovered, DNA match, tamper notes."],
            ["19–20 Certificates", "Need portable authenticity proof.", "Issue PDF + public verify ACTIVE."],
            ["25–27 Business shell", "Company cannot share one personal login.", "Business dashboard + org setup complete."],
            ["29 Team", "Team needs roles, not one shared account.", "Invite by PinIT ID; role shown."],
            ["30 Business protect", "Business must still protect files.", "Protect → Digital Assets → Investigate works."],
            ["33 Business E2E", "Need one company walkthrough.", "Setup → Team → Protect → Assets → Dashboard."],
        ],
        col_widths=[1.4, 2.7, 2.7],
    )

    # ── 7. Proud solutions demo ──
    add_heading(doc, "7. Recommended Proud Solutions Demo (One Hero Video)", 1)
    add_p(
        doc,
        "After individual test videos, record one continuous demo (~6–8 minutes) that tells the full story. "
        "This is the video leadership can share proudly.",
    )
    add_table(
        doc,
        [
            ["Scene", "Minutes", "What to show"],
            ["1. Login + Dashboard", "0–1", "Secure entry and Hub overview"],
            ["2. Generate DNA", "1–2", "Upload file → DNA complete → note DNA ID"],
            ["3. Vault", "2–3", "File in Vault Explorer → encryption visible"],
            ["4. Secure Share", "3–4.5", "Create link → open in Incognito → allow GPS → view"],
            ["5. Tracking proof", "4.5–5.5", "Timeline event + Access Intelligence map pin"],
            ["6. Investigation / Certificate", "5.5–7", "Intelligence Report or Certificate verify"],
            ["7. Close", "7–8", "“PinIT Hub protects the full file lifecycle.” End on Dashboard"],
        ],
        col_widths=[2.0, 1.0, 3.8],
    )
    add_p(doc, "Save as: PINIT_ProudSolutions_Demo_YYYY-MM-DD.mp4", bold=True)

    add_heading(doc, "7.1 Business Account videos (Part F — new)", 2)
    add_p(
        doc,
        "The original User Test Flow was written before Business account shipped. "
        "Part F (Tests 25–33) is required for end-stage testing. Use a Business test account.",
    )
    add_table(
        doc,
        [
            ["Test", "Video name", "Say this"],
            ["25", "PINIT_Test25_BusinessShell_…", "Business shell vs personal dashboard"],
            ["26", "PINIT_Test26_BusinessDashboard_…", "Org ops home and snapshots"],
            ["27", "PINIT_Test27_OrgSetupWizard_…", "Organization name and workspace setup"],
            ["28", "PINIT_Test28_OrgProfile_…", "Organization Profile hub tabs"],
            ["29", "PINIT_Test29_TeamRoles_…", "Invite teammate by PinIT ID"],
            ["30", "PINIT_Test30_BusinessProtect_…", "Same DNA vault stack inside Business"],
            ["33", "PINIT_Test33_BusinessE2E_…", "Full company mini walkthrough"],
        ],
        col_widths=[0.8, 2.8, 3.2],
    )

    # ── 8. Delivery tracker ──
    add_heading(doc, "8. Delivery Tracker (Fill While Testing)", 1)
    add_p(doc, "Copy this into your notes or tick in this document printout.")
    add_table(
        doc,
        [
            ["Video", "File name", "Result", "Notes"],
            ["Landing / Hub Demo", "PINIT_LandingDemo_…", "☐ Pass ☐ Fail", ""],
            ["Test 1 Secure Login", "PINIT_Test01_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Test 2 Dashboard", "PINIT_Test02_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Test 3 Generate DNA", "PINIT_Test03_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Test 4 Duplicate", "PINIT_Test04_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Tests 5–9 Vault/DNA/Timeline", "PINIT_Test05–09_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Tests 10–15 Share & Tracking", "PINIT_Test10–15_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Tests 16–20 Intel & Forensics", "PINIT_Test16–20_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Tests 21–24 Extra modules", "PINIT_Test21–24_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Tests 25–33 Business Account", "PINIT_Test25–33_…", "☐ Pass ☐ Fail ☐ Blocked", ""],
            ["Proud Solutions Demo", "PINIT_ProudSolutions_…", "☐ Done", ""],
        ],
        col_widths=[2.2, 2.0, 1.8, 0.8],
    )

    add_table(
        doc,
        [
            ["Field", "Fill in"],
            ["Tester name", ""],
            ["Date", ""],
            ["Environment", "Production — https://dna-pinit-web.vercel.app"],
            ["Browser", ""],
            ["Videos uploaded to", ""],
            ["Blockers reported to", ""],
        ],
        col_widths=[2.2, 4.6],
    )

    # ── 9. Quality rules ──
    add_heading(doc, "9. Quality Rules (Do / Don’t)", 1)
    add_heading(doc, "Do", 2)
    add_bullets(
        doc,
        [
            "Speak the problem and the proof.",
            "Move slowly so text and IDs are readable.",
            "Allow location when testing share viewer.",
            "Use Incognito for recipient role.",
            "Record again if a popup or notification ruined the clip.",
            "Keep one video per test (easier review).",
        ],
    )
    add_heading(doc, "Don’t", 2)
    add_bullets(
        doc,
        [
            "Don’t silently click with no explanation.",
            "Don’t show real personal identity documents of real people.",
            "Don’t skip expected-result checkmarks.",
            "Don’t mark Pass if the screen shows an error.",
            "Don’t combine all 24 tests into one unreadable mega-file (except the optional Proud Solutions demo).",
        ],
    )

    # ── 10. Quick start ──
    add_heading(doc, "10. Quick Start — Your First Hour", 1)
    add_numbered(
        doc,
        [
            "Read this guide fully (15 minutes).",
            "Open User Test Flow PDF next to you.",
            "Set up recording + test files + credentials (10 minutes).",
            "Record Landing / Hub Demo video (Section 4).",
            "Record Test Case 1 Secure Login (Section 5).",
            "Continue Tests 2–3 the same day if time allows.",
            "Upload videos + send Pass/Fail summary to the team.",
        ],
    )

    add_hline(doc)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run(
        "PinIT Hub — Secure · Connect · Control\n"
        "Walkthrough videos should make every viewer feel: we can prove ownership, protect files, "
        "track sharing, and investigate leaks.\n"
        "Use with: PINIT-DNA Complete User Test Flow & Application Guide"
    )
    set_run_font(run, size=10, color=DARK)

    try:
        doc.save(OUT_PATH)
        print(f"Wrote {OUT_PATH}")
    except PermissionError:
        doc.save(ALT_OUT_PATH)
        print(f"Main file locked — wrote {ALT_OUT_PATH}")


if __name__ == "__main__":
    build()
