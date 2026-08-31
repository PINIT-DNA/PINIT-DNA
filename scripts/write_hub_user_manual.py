"""Generate docs/Pinit_HUB_User_Manual_Walkthrough.docx"""
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "Pinit_HUB_User_Manual_Walkthrough.docx"


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_doc(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for i, size in ((1, 22), (2, 16), (3, 13)):
        s = styles[f"Heading {i}"]
        s.font.name = "Calibri"
        s.font.color.rgb = RGBColor(15, 23, 42)
        s.font.size = Pt(size)
        s.font.bold = True


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=26, bold=True, color=(15, 23, 42))
    p.paragraph_format.space_after = Pt(4)


def add_meta(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=10, color=(71, 85, 105))
    p.paragraph_format.space_after = Pt(16)


def p(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(8)
    return para


def bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    return para


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def step(doc, n, title, body_lines):
    h = doc.add_paragraph()
    r = h.add_run(f"Step {n}. {title}")
    set_run_font(r, size=12, bold=True, color=(30, 64, 175))
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for line in body_lines:
        p(doc, line)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
            for paragraph in t.rows[ri + 1].cells[ci].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    style_doc(doc)

    add_title(doc, "Pinit HUB — User Manual Walkthrough")
    add_meta(
        doc,
        "For Individual accounts and Business accounts  ·  Current product as of August 2026  ·  "
        "Live app: https://pinit-dna.vercel.app",
    )
    p(
        doc,
        "This is a click-by-click walkthrough of Pinit HUB as it works today. "
        "It is written for people using the product — not for engineers. "
        "Pinit Exchange (the marketplace at pinitexchange.com) is a separate app; this manual covers HUB only, "
        "and notes where HUB can open Exchange.",
    )

    doc.add_heading("1. What Pinit HUB is for", level=1)
    p(
        doc,
        "Pinit HUB protects original files (photos, videos, documents) so you still have proof and control after they leave your laptop. "
        "A protected file gets a DNA fingerprint, sits in an encrypted vault, can be shared with a tracked link, can be watched for leaks, and can be investigated if a copy shows up elsewhere.",
    )
    p(doc, "There are two ways to use HUB:")
    bullet(doc, "Individual account — you work as yourself. Home is the personal dashboard.")
    bullet(doc, "Business account — you work as a company with clients, campaigns, and a team. Home is the organization dashboard.")
    p(
        doc,
        "Both account types use the same protect / vault / share / investigate stack. Business adds organization, clients, campaigns, people, and roles on top.",
    )

    doc.add_heading("2. Where to open HUB", level=1)
    table(
        doc,
        ["Place", "Address"],
        [
            ["Live HUB", "https://pinit-dna.vercel.app"],
            ["Local development", "http://localhost:3000"],
            ["Sign in", "/login"],
            ["Create account", "/register/account-type then /register"],
            ["Pinit Exchange (marketplace)", "https://www.pinitexchange.com — opened from HUB when allowed"],
        ],
    )
    p(doc, "Your login identity is your Pinit ID (for example PINIT-XXXXXX), plus face authentication. Keep that Pinit ID — teammates invite you with it.")

    doc.add_heading("3. Create an account (first-time walkthrough)", level=1)

    step(
        doc,
        1,
        "Choose Individual or Business before you register",
        [
            "Open HUB and go to Create account. You land on Account type.",
            "Pick Individual if you work alone (photographer, freelancer, consultant).",
            "Pick Business if you run or join a company workspace (agency, studio, brand team).",
            "This choice is stored before biometrics so registration already knows which home to open.",
            "Click Continue. You go to registration.",
        ],
    )
    step(
        doc,
        2,
        "Complete registration (name, biometrics, Pinit ID)",
        [
            "Follow the on-screen registration flow. HUB uses face capture (liveness) so the account is bound to a real person.",
            "When registration finishes you receive a Pinit ID. Write it down. Other people search and invite you by this ID, not by exposing your email in campaign invites.",
            "If you already had a session, you will not re-register — you will be sent to Home.",
        ],
    )
    step(
        doc,
        3,
        "Confirm account type if HUB asks again",
        [
            "If account type was not saved, HUB opens Onboarding → Account type.",
            "Choose Individual or Business and continue.",
            "Individual opens the personal Home (/).",
            "Business opens the organization Home (/business) and may show a setup wizard.",
        ],
    )

    doc.add_heading("4. Sign in (every later visit)", level=1)
    numbered(doc, "Open https://pinit-dna.vercel.app (or localhost:3000).")
    numbered(doc, "Use Face login. HUB verifies it is you — it does not skip biometrics just because a leftover token exists.")
    numbered(doc, "You land on Individual Home or Business Home, depending on your account.")
    p(
        doc,
        "If someone sent you a team invitation link (/team/join/…) while you were logged out, HUB remembers that token. After you sign in as the correct Pinit account, acceptance continues. Opening the link alone does not grant access.",
    )

    doc.add_heading("5. Find your way around", level=1)
    doc.add_heading("5.1 Individual sidebar", level=2)
    table(
        doc,
        ["Menu", "What it is"],
        [
            ["Home", "Personal dashboard — counts, recent files, security insights"],
            ["Protect asset", "Upload a file and generate DNA + vault record"],
            ["Digital Assets", "Your vault — originals, share, retrieve"],
            ["Protected assets", "DNA records list"],
            ["Monitoring", "Watch for copies appearing on the web (plan-gated)"],
            ["Investigate", "Upload a leaked/probe file and match it to your vault"],
            ["Reports", "Investigation / forensic reports"],
            ["Unmask Requests", "Review requests to see masked content"],
            ["Duplicate Attempts", "Blocked or duplicate registration/protect attempts"],
            ["Vault check", "Integrity check on vault files"],
            ["Certificates", "Issue and list authenticity certificates"],
            ["Verify Certificate", "Check a certificate without owning the file"],
            ["Plans & Upgrade / Subscription", "Free / Pro (and billing)"],
        ],
    )
    doc.add_heading("5.2 Business sidebar", level=2)
    table(
        doc,
        ["Menu", "What it is"],
        [
            ["Home", "Organization operations dashboard"],
            ["Clients", "Companies you work for — then campaigns under each client"],
            ["Protect asset", "Same protect flow; you can attach the file to a campaign"],
            ["Digital Assets", "Vault (your authenticated identity still owns files)"],
            ["Investigate / Reports / Monitoring / Certificates", "Same protection stack as Individual"],
            ["Team", "Organization members and invitations (via Profile team tab)"],
            ["Audit Logs", "Who did what in the organization"],
            ["API Access", "Keys for integrations (Enterprise-gated)"],
            ["Organization Profile", "Company name, settings, billing snapshot"],
        ],
    )
    doc.add_heading("5.3 Switch Individual / Business / Exchange", level=2)
    p(
        doc,
        "If your account is allowed to switch views, the top of the dashboard shows Individual | Business | Exchange. "
        "Individual and Business change which HUB home you see. Exchange opens the marketplace in a signed-in session as your same Pinit ID — it is not a second HUB account.",
    )

    doc.add_heading("6. Individual account — full walkthrough", level=1)
    p(doc, "Use this path if you registered as Individual, or switched to Individual view.")

    doc.add_heading("6.1 Home", level=2)
    numbered(doc, "After login, you are on Home (/).")
    numbered(doc, "Read the stat cards (protected files, vault size, shares, risk). Click a card to jump to that area.")
    numbered(doc, "Scroll security insights (views, blocked downloads, duplicate attempts).")
    numbered(doc, "Use Quick actions or Protect asset when you are ready to fingerprint a file.")

    doc.add_heading("6.2 Protect a file", level=2)
    numbered(doc, "Sidebar → Protect asset (/generate).")
    numbered(doc, "Upload an image, video, or document.")
    numbered(doc, "Wait for DNA generation (multi-layer fingerprint) and vault encryption.")
    numbered(doc, "When it finishes, the file exists as a DNA record and as an encrypted vault original. You have not “posted” it anywhere yet — it is still only in HUB.")
    p(doc, "Tip: Protect first, then share. Sharing an unprotected file from chat or Drive is outside HUB and will not create a tracked trail.")

    doc.add_heading("6.3 Digital Assets (vault)", level=2)
    numbered(doc, "Sidebar → Digital Assets (/vault).")
    numbered(doc, "Open a file to see type, size, dates, and actions.")
    numbered(doc, "Retrieve / download uses HUB’s protected download path where enabled — it is not a silent public file dump.")
    numbered(doc, "To share: open Share on that asset (or /vault/assets/{id}/share).")

    doc.add_heading("6.4 Share with a tracked link", level=2)
    numbered(doc, "Create a Smart Share / tracked link for a vault asset.")
    numbered(doc, "Copy the link. Recipients open /s/{token} — they do not need a HUB login for a normal share.")
    numbered(doc, "You (the owner) later open Access Intelligence or the share manage page to see views, hops, and risk.")
    numbered(doc, "A person who only received a share link does not become a Business team member or a campaign person. Share access is external.")
    p(doc, "If you used masking, the recipient may request an unmask. You review that under Unmask Requests.")

    doc.add_heading("6.5 Access Intelligence", level=2)
    numbered(doc, "Open Access Intelligence from Home insights, a share, or /access-intelligence.")
    numbered(doc, "Inspect who opened the link, from where, and whether the file was forwarded (link tree / hop chain when available).")

    doc.add_heading("6.6 Monitor the web", level=2)
    numbered(doc, "Sidebar → Monitoring (requires tracking feature on your plan).")
    numbered(doc, "Enroll an asset if you want HUB to watch for public copies.")
    numbered(doc, "Alerts appear when the crawler/providers are configured. Empty monitoring is normal on a fresh account.")

    doc.add_heading("6.7 Investigate a leak or lookalike", level=2)
    numbered(doc, "Sidebar → Investigate (/pinit-hub/investigation).")
    numbered(doc, "Upload the leaked or suspicious file (the “probe”).")
    numbered(doc, "HUB compares it to your vault DNA and shows match / tamper / forensic results.")
    numbered(doc, "Open Reports to keep or export investigation output.")
    numbered(doc, "Use Forensic Diff when you need a visual comparison of two versions.")

    doc.add_heading("6.8 Certificates", level=2)
    numbered(doc, "Sidebar → Certificates to issue or list certificates for a protected asset.")
    numbered(doc, "Anyone can open Verify Certificate to check a certificate without logging in as you.")

    doc.add_heading("6.9 Plans", level=2)
    numbered(doc, "Plans & Upgrade or Subscription to see Free vs Pro, storage, and team limits.")
    numbered(doc, "Checkout / payment pages complete an upgrade. After success you may see a welcome modal once.")

    doc.add_heading("6.10 Profile and security", level=2)
    numbered(doc, "Open Profile (avatar menu).")
    numbered(doc, "Profile — display name (this is what teammates see, not your email, on invites).")
    numbered(doc, "Security — face / passkey related controls.")
    numbered(doc, "Notifications — in-app alerts (invites, shares, investigations).")
    numbered(doc, "Activity — what your account did recently.")
    numbered(doc, "Settings — theme and preferences.")

    doc.add_heading("7. Business account — full walkthrough", level=1)
    p(doc, "Use this path if you registered as Business, or switched to Business view.")

    doc.add_heading("7.1 First-time organization setup", level=2)
    numbered(doc, "You land on Business Home (/business).")
    numbered(doc, "If setup is incomplete, complete the Business Setup Wizard: organization name, industry, country, workspace, optional logo.")
    numbered(doc, "You become the organization Owner. That is a business/organization role — it is not the same as a campaign role.")

    doc.add_heading("7.2 Business Home", level=2)
    p(doc, "The operations dashboard summarizes clients, vault, team, investigations, monitoring, certificates, alerts, and recent activity. Use Quick actions to add a client, protect a file, or invite a teammate.")

    doc.add_heading("7.3 Clients", level=2)
    numbered(doc, "Sidebar → Clients (/business/clients).")
    numbered(doc, "Click Add client. Enter the client (brand) name and optional contact.")
    numbered(doc, "Open a client. This is the Client workspace.")
    p(doc, "Client tabs:")
    bullet(doc, "Campaigns — jobs for this client. Create a campaign here.")
    bullet(doc, "Assets — files attached to this client’s campaigns.")
    bullet(doc, "People — people appearing across those campaigns.")
    bullet(doc, "Deliveries, Rights, Activity, Intelligence — rollups of handover, licensing, history, and insights.")
    p(doc, "The client themselves does not log into your Business Account. They review through secure links you send (share, handover, or client report).")

    doc.add_heading("7.4 Create a campaign", level=2)
    numbered(doc, "Inside a client, open Campaigns → create campaign.")
    numbered(doc, "Name it (for example “Summer launch”), optional dates and notes.")
    numbered(doc, "Open the campaign. You are in the Campaign workspace.")

    doc.add_heading("7.5 Campaign workspace tabs", level=2)
    table(
        doc,
        ["Tab", "What you do here"],
        [
            ["Overview", "Status, counts, what needs attention (open change requests, unread messages)"],
            ["Assets", "Files protected into this campaign"],
            ["People", "Internal team vs external creators; add person"],
            ["Approvals", "Change requests and review decisions"],
            ["Versions", "Version history of deliverables"],
            ["Messages", "Campaign discussion"],
            ["Rights", "Usage rights / licensing notes"],
            ["Handover", "Package and send a handover (public /handover/{token})"],
            ["Monitoring", "Campaign-scoped monitoring"],
            ["Findings", "Issues found on campaign assets"],
            ["Investigations", "Investigations tied to this campaign"],
            ["Sharing", "Tracked share links for campaign assets"],
            ["Activity", "Audit-style campaign activity"],
            ["Intelligence", "Client/campaign intelligence snapshot"],
        ],
    )

    doc.add_heading("7.6 Protect a file into a campaign", level=2)
    numbered(doc, "From the campaign, use Protect (or sidebar Protect asset with the campaign selected when the flow offers campaignId).")
    numbered(doc, "The asset stays one canonical file in the vault. The campaign only references it — HUB does not duplicate the file into a second campaign copy.")
    numbered(doc, "After protect, the file should appear under Campaign → Assets and Digital Assets.")

    doc.add_heading("7.7 People — add a team member (organization colleague)", level=2)
    numbered(doc, "Campaign → People → Add team member (or Add person to campaign).")
    numbered(doc, "Keep Team member selected (“Someone in your organization”). Do not use External creator for staff.")
    p(doc, "Path A — they already belong to your organization:")
    numbered(doc, "Choose Existing organization member.")
    numbered(doc, "Search by name or Pinit ID.")
    numbered(doc, "Select the person. You should see their display name and Pinit ID.")
    numbered(doc, "Pick a Campaign role: Campaign owner, Project manager, Contributor, Reviewer, Designer, or Developer.")
    numbered(doc, "Click Add to campaign.")
    p(
        doc,
        "They are added immediately as a CampaignMember bound to their real userId. Their Business/Organization role does not change. Example: Business role can stay MEMBER while Campaign role is Designer. The same campaign appears in their own HUB when they sign in.",
    )
    p(doc, "Path B — they have a Pinit account but are not in your organization yet:")
    numbered(doc, "Choose Invite by Pinit ID.")
    numbered(doc, "Enter PINIT-XXXXXX → Find account.")
    numbered(doc, "If nobody exists: “No Pinit account with that ID”.")
    numbered(doc, "If found: you see verified name + Pinit ID only (not email, not internal user id).")
    numbered(doc, "If they are already on the campaign, stop — do not duplicate.")
    numbered(doc, "If they are already in the org, Add to campaign (same as Path A).")
    numbered(doc, "Otherwise pick Campaign role → Send invitation.")
    numbered(doc, "Copy the /team/join/{token} link. Opening it does not grant access.")
    numbered(doc, "People tab shows Invitation pending until they accept.")
    p(doc, "Recipient walkthrough:")
    numbered(doc, "They open the link while signed in as that exact Pinit ID.")
    numbered(doc, "Wrong account: “This invitation is for a different Pinit account”.")
    numbered(doc, "Expired or already used invites are rejected.")
    numbered(doc, "On accept they join the organization (if needed) and this same campaign. They see Campaign A in their HUB — not a copy.")

    doc.add_heading("7.8 People — add an external creator", level=2)
    numbered(doc, "In Add person, choose External creator (“Freelancer or influencer”).")
    numbered(doc, "Enter name, optional platform and profile URL, and a campaign role label.")
    numbered(doc, "They do not become an organization member.")
    numbered(doc, "Being on the People list is not asset access. Expand the person and assign specific assets + permissions, then share tracked links.")
    p(doc, "A normal asset share recipient is also not an internal campaign person unless you add them here on purpose.")

    doc.add_heading("7.9 Organization Team (company membership)", level=2)
    numbered(doc, "Open Team from the business sidebar or Profile team tab (/business/team or /profile?tab=team).")
    numbered(doc, "Invite by Pinit ID for company membership (Owner cannot be invited; typical invite role is Member).")
    numbered(doc, "Organization roles (Owner, Manager, Member, Viewer, …) control who can add people, invite, and change org settings.")
    numbered(doc, "A normal Member who can open a campaign page still cannot add members unless they have Manager (or higher) organization permission. Campaign roles do not replace that.")

    doc.add_heading("7.10 Review, messages, handover, client report", level=2)
    numbered(doc, "Approvals — request or resolve changes on versions.")
    numbered(doc, "Messages — campaign thread; unread shows on the tab.")
    numbered(doc, "Handover — generate a handover package/link. Recipients use /handover/{token}.")
    numbered(doc, "Client report — public /client-report/{token} for a client-facing summary when you issue one.")
    numbered(doc, "Sharing — use existing tracked share links; every view/download is logged.")

    doc.add_heading("7.11 Organization Profile, audit, API", level=2)
    numbered(doc, "Organization Profile — company details and operating settings.")
    numbered(doc, "Audit Logs — organization actions (invites, campaign member added, and similar).")
    numbered(doc, "API Access — Enterprise. Do not create keys on a plan that does not include API.")

    doc.add_heading("8. Two kinds of “role” (do not mix them)", level=1)
    table(
        doc,
        ["Kind", "Examples", "What it controls"],
        [
            ["Organization / business role", "Owner, Manager, Member, Viewer", "Who can invite, add campaign people, change org settings"],
            ["Campaign role", "Campaign owner, Project manager, Contributor, Reviewer, Designer, Developer", "Job on that campaign only — does not change business role"],
        ],
    )
    p(doc, "Making someone Campaign Owner does not make them Organization Owner.")

    doc.add_heading("9. Public links (no HUB login required, except team join)", level=1)
    table(
        doc,
        ["Link", "Who uses it", "What happens"],
        [
            ["/s/{token}", "Share recipient", "View/download per share rules; logged; not org membership"],
            ["/team/join/{token}", "Invited Pinit user", "Must sign in as the invited Pinit ID and accept"],
            ["/handover/{token}", "Client / recipient of a handover", "Sees the handover package"],
            ["/client-report/{token}", "Client", "Sees the issued client report"],
            ["/verify-certificate", "Anyone", "Checks a certificate"],
        ],
    )

    doc.add_heading("10. Suggested first day (Individual)", level=1)
    numbered(doc, "Register as Individual and save your Pinit ID.")
    numbered(doc, "Protect one photo.")
    numbered(doc, "Open it in Digital Assets.")
    numbered(doc, "Create a tracked share to yourself in another browser and open it.")
    numbered(doc, "Open Access Intelligence and confirm the view.")
    numbered(doc, "Optional: run Investigate with a cropped copy of the same photo.")

    doc.add_heading("11. Suggested first day (Business)", level=1)
    numbered(doc, "Register as Business, complete org setup, save your Pinit ID.")
    numbered(doc, "Add a client.")
    numbered(doc, "Create Campaign A under that client.")
    numbered(doc, "Protect one file into Campaign A.")
    numbered(doc, "People → add an existing teammate, or Invite by Pinit ID, assign Contributor.")
    numbered(doc, "Have them sign in on their own HUB and open the same Campaign A.")
    numbered(doc, "Add one External creator without inviting them to the organization.")
    numbered(doc, "Share one asset with a tracked link and confirm the recipient is not added as staff.")

    doc.add_heading("12. Common problems", level=1)
    table(
        doc,
        ["What you see", "What to do"],
        [
            ["No Pinit account with that ID", "The ID is wrong or unused. Ask the person for the ID shown in their Profile."],
            ["This invitation is for a different Pinit account", "Sign out and sign in as the invited Pinit ID. Do not try to accept as someone else."],
            ["Invitation pending", "They have not accepted yet. Resend the same link; do not stack extra invites."],
            ["Already connected to this campaign", "They are already a CampaignMember. Do not add again."],
            ["Cannot add person / forbidden", "Your organization role is not Manager+. Campaign access is not enough."],
            ["Share opened but they are not on People", "Correct. Share recipients stay external until you add them as team or external creator."],
            ["They accepted but I expected a second campaign", "HUB never copies the campaign. They see the same Campaign A."],
        ],
    )

    doc.add_heading("13. What this manual does not cover", level=1)
    bullet(doc, "Pinit Exchange buying/selling — use Exchange after opening it from HUB.")
    bullet(doc, "Super-admin console (/admin) — internal operators only.")
    bullet(doc, "Chrome extension / Publish Guardian deep setup — separate extension guide.")
    bullet(doc, "Server, database, or deploy operations.")

    doc.add_heading("14. Quick map of both lives in one picture", level=1)
    p(doc, "Individual: Register → Home → Protect → Vault → Share → Access Intelligence → Monitor → Investigate → Certificates → Profile.")
    p(doc, "Business: Register → Org setup → Business Home → Client → Campaign → Protect into campaign → People (internal invite or external creator) → Review / Messages / Handover / Sharing → Team & Audit. Protection, vault, investigate, and certificates stay available the whole time.")
    p(
        doc,
        "Remember: one person, one Pinit ID, one vault identity. Business is an organization layer around that identity. Campaigns are jobs under a client. Shares are how outsiders view files without joining the company.",
    )

    footer = doc.add_paragraph()
    r = footer.add_run("Pinit HUB User Manual Walkthrough  ·  Individual & Business  ·  August 2026")
    set_run_font(r, size=9, color=(100, 116, 139))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
