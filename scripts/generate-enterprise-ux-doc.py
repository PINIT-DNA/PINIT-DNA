"""Generate PINITHub Enterprise UX Design Spec as Word document."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "enterprise" / "PINITHub_Enterprise_UX_Design_Spec.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "1E293B") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PINITHub Enterprise", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Complete UX & Product Design Specification")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph("Version 1.0  |  Pre-implementation validation  |  July 2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Reference scenario: ABC Media Pvt Ltd purchases Enterprise (₹4,999/month)")
    doc.add_page_break()

    doc.add_heading("Design North Star", 1)
    doc.add_paragraph(
        "PINITHub Enterprise should feel like GitHub Enterprise meets Microsoft Purview — "
        "governance-first, collaboration-native, audit-ready. Every screen answers:"
    )
    add_bullets(doc, [
        "Who owns this asset?",
        "Who can access or investigate it?",
        "What happened, and can we prove it?",
    ])

    doc.add_heading("Visual Language", 2)
    add_table(doc,
        ["Token", "Enterprise usage"],
        [
            ["Dark shell", "Primary workspace — reduces eye strain for long investigations"],
            ["DNA accent", "Actions, links, active navigation"],
            ["Amber", "Pro-tier indicators, warnings"],
            ["Purple", "Enterprise-only features, org admin"],
            ["Emerald", "Verified / safe / approved"],
            ["Red", "Threat / leak / critical alert"],
            ["Typography", "Inter/system — 13px body, 11px meta, 18–22px section titles"],
            ["Density", "Compact tables (GitHub-style), generous cards on dashboard"],
        ],
    )

    doc.add_heading("Global Navigation Architecture", 1)
    doc.add_heading("Context Switcher (top-left)", 2)
    add_code_block(doc, "[ ABC Media logo ]  ABC Media Pvt Ltd  ▾\n  └── Switch organization\n  └── Personal workspace\n  └── Create organization")

    doc.add_heading("Enterprise Sidebar (240px)", 2)
    add_code_block(doc, """ORGANIZATION
  Dashboard | Shared Vault | Investigations | Monitoring | Reports
TEAM
  Members | Teams | Invitations
GOVERNANCE
  Audit Logs | Compliance | Policies
ADMIN (Owner/Admin only)
  Settings | Billing | API & Integrations
───────────────
[ User avatar ]  Priya Sharma · Owner
  Enterprise · Unlimited storage""")

    doc.add_heading("Top Bar", 2)
    add_code_block(doc, "Breadcrumb: ABC Media > Shared Vault > Marketing > Q3 Campaign\n[ Search ⌘K ]  [ + Generate DNA ]  [ Notifications ]  [ Avatar ]")

    # Section 1
    doc.add_page_break()
    doc.add_heading("1. Enterprise Purchase Flow", 1)
    doc.add_heading("Screen 1 — Upgrade Page (enhanced)", 2)
    add_code_block(doc, """Choose Enterprise for your organization
Unlimited storage · Teams · API · Audit · Monitoring

Free ₹0  |  Pro ₹999/mo  |  ★ Enterprise ₹4,999/mo
                              [ Pay with Razorpay ]""")
    doc.add_paragraph("Enterprise copy: For companies with 5+ users, shared vaults, and compliance requirements.")

    doc.add_heading("Screen 2 — Payment Processing", 2)
    doc.add_paragraph("Full-screen Razorpay modal. Background shows blurred upgrade page.")

    doc.add_heading("Screen 3 — Payment Success → Organization Bootstrap", 2)
    add_code_block(doc, """✓ Payment successful
Welcome to PINITHub Enterprise
Let's set up ABC Media Pvt Ltd

What you'll configure:
○ Organization profile
○ Team workspace
○ Invite your first members

[ Continue setup → ]  Skip for now""")
    doc.add_paragraph("System actions: Organization created, purchaser linked as Owner, default workspace Main created, Enterprise entitlements activated, audit event SUBSCRIPTION_ENTERPRISE_ACTIVATED.")

    doc.add_heading("Screen 4 — Welcome Wizard", 2)
    doc.add_paragraph("Progress bar: Step 1 of 4 through Step 4 of 4.")

    # Section 2
    doc.add_heading("2. Organization Setup", 1)
    doc.add_heading("Step 1 — Organization Profile", 2)
    add_code_block(doc, """Organization name *     [ ABC Media Pvt Ltd ]
Logo (optional)         Industry *    [ Media & Entertainment ▾ ]
Company size *          [ 51–200 ▾ ]  Country * [ India ▾ ]
Primary domain          [ abmedia.in ]""")

    doc.add_heading("Step 2 — Create Workspace", 2)
    add_code_block(doc, """Workspace name *        [ Main Operations ]
Description             [ Central vault for marketing, legal, production ]
Default departments: ☑ Marketing ☑ Legal ☑ Production ☐ HR ☐ Finance""")

    doc.add_heading("Step 3 — Invite Team", 2)
    add_code_block(doc, """Email addresses (comma separated)
[ priya@abmedia.in, legal@abmedia.in, ops@abmedia.in ]
Default role: [ Team Member ▾ ]
☐ Send welcome email
[ Skip ]  [ Send invitations → ]""")

    doc.add_heading("Step 4 — Ready", 2)
    add_code_block(doc, """✓ ABC Media Pvt Ltd is ready
Organization: ABC Media Pvt Ltd | Workspace: Main Operations
Invites sent: 3 pending | Storage: Unlimited
[ Go to Enterprise Dashboard → ]""")

    # Section 3
    doc.add_page_break()
    doc.add_heading("3. Team Workspace Layout", 1)
    doc.add_heading("Enterprise Dashboard", 2)
    add_code_block(doc, """Good morning, Priya · ABC Media Pvt Ltd                    [This week ▾]

Organization Assets 1,247  |  DNA Generated 342  |  Storage 847 GB / ∞
Active Investigations (4)  |  Threat Alerts (2 critical)
Recent Activity            |  Team (12 members, +3 pending)
Subscription: Enterprise · Renews Apr 21 · 12 seats · API: 24.1K calls""")

    doc.add_heading("Shared Vault Layout", 2)
    add_code_block(doc, """FOLDERS          |  Shared Vault > Marketing > Q3 Campaign  [+ Upload]
All Assets       |  [ Search ] [ Type ▾ ] [ Owner ▾ ] [ Grid|List ]
Marketing        |  □ hero_banner.jpg   DNA ✓  Priya   Mar 12  Shared
  Q3 Campaign    |  □ promo_video.mp4   DNA ✓  Rahul   Mar 10  Restricted
Legal            |  Right panel: Preview · DNA · Permissions · Log""")

    # Section 4
    doc.add_heading("4. Invite Members Flow", 1)
    add_code_block(doc, """Invite members to ABC Media
Email [ user@... ]  Role [ Investigator ▾ ]  Team [ Legal ▾ ]
Personal message (optional)
[ Cancel ]  [ Send invitation ]""")
    add_table(doc,
        ["Email", "Role", "Invited by", "Sent", "Status", "Actions"],
        [
            ["legal@abmedia.in", "Investigator", "Priya", "Mar 10", "Pending", "Resend · Revoke"],
            ["ops@abmedia.in", "Team Member", "Priya", "Mar 10", "Expired", "Resend"],
        ],
    )
    doc.add_paragraph("Accept invitation lands on /invite/{token} → biometric login → role assigned → workspace.")

    # Section 5
    doc.add_page_break()
    doc.add_heading("5. User Roles & Permission Matrix", 1)
    add_table(doc,
        ["Role", "Persona", "Primary use"],
        [
            ["Owner", "CTO / Founder", "Billing, org delete, all permissions"],
            ["Admin", "IT / Security lead", "Members, roles, settings, API keys"],
            ["Manager", "Dept head", "Team vaults, assign investigations, approve shares"],
            ["Investigator", "Legal / Forensics", "Run investigations, view evidence, no billing"],
            ["Team Member", "Creator", "Upload, DNA, store, share within policy"],
            ["Viewer", "External counsel", "Read-only vault + reports"],
        ],
    )

    doc.add_heading("Permission Matrix", 2)
    add_table(doc,
        ["Capability", "Owner", "Admin", "Manager", "Investigator", "Member", "Viewer"],
        [
            ["Manage billing", "✓", "✓", "—", "—", "—", "—"],
            ["Delete organization", "✓", "—", "—", "—", "—", "—"],
            ["Manage members/roles", "✓", "✓", "—", "—", "—", "—"],
            ["API keys & webhooks", "✓", "✓", "—", "—", "—", "—"],
            ["Upload & generate DNA", "✓", "✓", "✓", "✓", "✓", "—"],
            ["Create Smart Share links", "✓", "✓", "✓", "—", "✓*", "—"],
            ["Run investigations", "✓", "✓", "✓", "✓", "—", "—"],
            ["Approve investigations", "✓", "✓", "✓", "—", "—", "—"],
            ["View audit logs", "✓", "✓", "✓", "✓", "—", "✓"],
            ["Monitoring & leak alerts", "✓", "✓", "✓", "✓", "—", "—"],
            ["Delete assets", "✓", "✓", "✓", "—", "—", "—"],
        ],
    )
    doc.add_paragraph("*Member share creation can require Manager approval via org policy.")

    doc.add_heading("Permission Flow", 2)
    add_bullets(doc, [
        "User action → Authenticated? → Org member? → Role permits? → Resource ACL? → Org policy? → Execute + audit log",
        "403 responses include suggest-request-access for denied actions.",
    ])

    # Section 6
    doc.add_heading("6. Shared Enterprise Vault", 1)
    add_code_block(doc, """Organization
 └── Workspace (Main Operations)
      └── Department (Marketing)
           └── Project (Q3 Campaign)
                └── Assets (files)""")
    doc.add_paragraph("Asset detail panel tabs: Overview | DNA | Access | Log. Actions: Share, Investigate, Download protected.")
    doc.add_paragraph("Search (⌘K): filename, DNA ID, tags, investigator notes. Filters: type, department, owner, date, share status.")

    # Section 7
    doc.add_page_break()
    doc.add_heading("7. Enterprise Dashboard Widgets", 1)
    add_table(doc,
        ["Widget", "Content", "Drill-down"],
        [
            ["Organization Assets", "Total count, trend", "Shared Vault"],
            ["DNA Generated", "30d count, success rate", "Reports"],
            ["Storage Usage", "Used / limit, by dept", "Billing > Usage"],
            ["Active Investigations", "Count, status breakdown", "Investigations list"],
            ["Threat Alerts", "Critical / warning", "Monitoring queue"],
            ["Monitoring Status", "Crawler health, last scan", "Monitoring"],
            ["Recent Activity", "Upload, share, investigate feed", "Audit log"],
            ["Team Members", "Count, pending invites", "Members"],
            ["Subscription Status", "Plan, renewal, seats", "Billing"],
            ["Compliance Score", "% assets with DNA + audit", "Compliance report"],
            ["API Usage", "Calls / rate limit", "API dashboard"],
            ["Investigation Queue", "Unassigned items", "Assign modal"],
        ],
    )

    # Section 8
    doc.add_heading("8. Team Asset Workflow", 1)
    add_bullets(doc, [
        "Team Member uploads → Generate DNA → Encrypt & Store",
        "Destination: Personal vault OR Shared Vault folder",
        "Team views asset → Suspicious activity? → Create Investigation OR Share/Monitor",
        "Investigation → Collect evidence → Acceptance verdict → Audit + Report",
        "Share path → Access Intelligence / TEP tracking → Report",
    ])
    doc.add_paragraph("Upload modal: Save to My vault / Shared Vault, folder picker, notify team on upload.")

    # Section 9
    doc.add_heading("9. Investigation Workflow (Multi-User)", 1)
    add_table(doc,
        ["ID", "Asset", "Status", "Owner", "Assignee", "Updated"],
        [
            ["INV-041", "leaked_ad.jpg", "In Progress", "Priya", "Rahul", "2h ago"],
            ["INV-040", "contract.pdf", "Review", "Ananya", "Legal", "1d ago"],
            ["INV-039", "video_clip.mp4", "Closed", "Priya", "Rahul", "3d ago"],
        ],
    )
    doc.add_paragraph("Investigation detail: 3-column layout — Evidence | Timeline & Comments | Verdict panel.")
    doc.add_paragraph("Status lifecycle: Draft → Assigned → In Progress → Pending Review → Approved → Closed")
    doc.add_paragraph("Manager or Owner must approve before AUTHORITATIVE verdict export to compliance.")

    # Section 10
    doc.add_page_break()
    doc.add_heading("10. Audit Logs", 1)
    add_table(doc,
        ["Time", "Actor", "Category", "Action", "Resource", "IP"],
        [
            ["11:02:03", "Rahul", "Upload", "VAULT_STORED", "contract", "103.*"],
            ["11:01:44", "System", "Investigation", "VERDICT_ISSUED", "INV-041", "—"],
            ["10:58:12", "Priya", "Role", "MEMBER_ROLE_CHANGED", "ops@...", "49.*"],
            ["10:45:00", "Ananya", "Share", "SHARE_LINK_CREATED", "hero.jpg", "49.*"],
        ],
    )
    add_table(doc,
        ["Category", "Events"],
        [
            ["Auth", "Login, logout, failed login, MFA, session revoke"],
            ["Assets", "Upload, download, delete, move, DNA generate"],
            ["Share", "Link created, revoked, accessed, OTP sent"],
            ["Investigation", "Created, assigned, comment, verdict, export"],
            ["Admin", "Role change, invite, policy change"],
            ["Billing", "Plan change, payment, invoice"],
            ["API", "Key created, webhook fired, rate limit hit"],
        ],
    )
    doc.add_paragraph("Properties: Immutable, append-only, 7-year retention (Enterprise), exportable for SOC2/ISO.")

    # Section 11
    doc.add_heading("11. Organization Settings", 1)
    add_bullets(doc, [
        "General — org name, logo, domain, industry",
        "Security — SSO (future), session timeout, IP allowlist, share approval",
        "Members & Roles — member list, role templates",
        "Teams — department teams",
        "Storage & Retention — quotas, retention policies",
        "Branding — logo, colors, custom domain on share viewer",
        "Billing — plan, invoices, payment method",
        "Notifications — org-wide alert preferences",
        "Integrations — SSO, Slack, webhooks",
        "Policies — max share duration, mandatory GPS, investigation approval",
        "Danger Zone — transfer ownership, downgrade, delete org",
    ])

    # Section 12
    doc.add_heading("12. Billing", 1)
    add_code_block(doc, """Current plan: Enterprise · ₹4,999/month · Renews Apr 21, 2026
Seats: 12 / unlimited · Storage: 847 GB / unlimited
Usage: Storage, API calls 24.1K/100K, Investigations 18, Share links 47
Invoices: Mar/Feb/Jan 2026 — ₹4,999 Paid [ PDF ]""")

    # Section 13
    doc.add_heading("13. API Access", 1)
    add_code_block(doc, """Tabs: API Keys | Webhooks | SDK | Usage | Logs
prod-vault-sync  Scopes: vault:read, dna:generate, investigation:read
Rate limits: 100 req/min · Webhooks: investigation.completed, leak.detected""")

    # Section 14
    doc.add_heading("14. Enterprise Monitoring", 1)
    doc.add_paragraph("Command center: Live geo map, investigation queue, alert rules, leak detection, crawler status.")

    # Section 15
    doc.add_heading("15. Notifications", 1)
    add_table(doc,
        ["Type", "Example", "Action"],
        [
            ["Threat", "Leak detected on asset X", "Open monitoring"],
            ["Investigation", "INV-041 assigned to you", "Open investigation"],
            ["Storage", "Org at 90% (if capped)", "Billing"],
            ["Team", "Member accepted invite", "Members"],
            ["Role", "Role changed to Investigator", "Profile"],
            ["Billing", "Invoice paid / failed", "Billing"],
            ["Monitoring", "Crawler completed — 2 matches", "Monitoring"],
        ],
    )

    # Section 16
    doc.add_heading("16. Enterprise Reports", 1)
    add_bullets(doc, [
        "Monthly Executive Summary (auto-generated)",
        "Asset Inventory Report",
        "Threat & Leak Report",
        "Investigation Summary",
        "Compliance & Chain of Custody",
        "Export PDF / CSV, schedule email to Owners",
    ])

    # Section 17
    doc.add_heading("17. Mobile Experience", 1)
    add_code_block(doc, "Bottom nav: [ Home ] [ Vault ] [ + ] [ Investigate ] [ More ]")
    add_bullets(doc, [
        "Home: condensed dashboard — alerts, recent activity",
        "Vault: folder drill-down, swipe share/investigate",
        "Investigate: camera/upload suspect file → quick DNA compare",
        "Admin tasks (billing, API, audit export): desktop-only banner",
        "Biometric login preserved; org context under user name",
    ])

    # Section 18
    doc.add_page_break()
    doc.add_heading("18. Complete User Journey", 1)
    add_bullets(doc, [
        "Company buys Enterprise → Razorpay payment → Org + Owner created → Welcome wizard",
        "Org profile → Create workspace → Invite team",
        "Member receives email → Accept + biometric login → Org dashboard",
        "Upload asset → Generate DNA → Store in Shared Vault",
        "Team views / shares → Issue detected → Create investigation",
        "Collaborate + evidence → Approve + close → Export report",
        "Audit log entry → Compliance report → Monthly executive review",
    ])

    doc.add_heading("Entity Relationship (Suggested Data Model)", 1)
    add_table(doc,
        ["Entity", "Relationships"],
        [
            ["Organization", "has Workspaces, OrgMembers, Subscription, AuditLogs, ApiKeys"],
            ["User", "belongs to OrgMember with Role"],
            ["Workspace", "contains Departments → Projects → VaultFolders → VaultRecords"],
            ["VaultRecord", "has DnaRecord, ShareLinks, Investigations"],
            ["Investigation", "has Comments, Evidence; assigned to User"],
            ["Team", "has TeamMembers; folder ACL via accessible_by"],
        ],
    )

    doc.add_heading("Suggested New Tables", 1)
    add_table(doc,
        ["Table", "Purpose"],
        [
            ["organizations", "Org profile, domain, branding"],
            ["org_members", "userId + orgId + roleId"],
            ["roles", "System + custom roles"],
            ["workspaces, departments, projects", "Folder hierarchy"],
            ["vault_folders", "Enterprise folder tree"],
            ["folder_permissions", "ACL per team/role/user"],
            ["investigations", "Case management"],
            ["investigation_comments, investigation_evidence", "Collaboration"],
            ["org_audit_logs", "Immutable audit stream"],
            ["org_invitations", "Pending invites"],
            ["api_keys, webhooks", "Developer access"],
        ],
    )
    doc.add_paragraph("Migration: extend vault records with optional organizationId + folderId; solo users keep implicit personal org.")

    doc.add_heading("UX Recommendations & Enterprise Best Practices", 1)
    doc.add_heading("Do", 2)
    add_bullets(doc, [
        "Progressive disclosure — Owners see billing; Investigators never do",
        "Context always visible — Org name + role in sidebar footer",
        "Audit by default — every destructive action confirms + logs",
        "@mentions in investigations — Slack-like collaboration",
        "Saved views for power users",
        "Bulk actions for hundreds of assets",
        "Export everything — PDF for legal, CSV for audit",
        "Empty states with clear next action",
    ])
    doc.add_heading("Avoid", 2)
    add_bullets(doc, [
        "Mixing personal and org assets without clear labels",
        "Hiding audit logs behind deep navigation",
        "Role escalation without Owner approval",
        "Mobile-only admin for billing/API",
        "Deleting audit logs (soft-delete only with legal hold)",
    ])

    doc.add_heading("Competitive Patterns to Borrow", 1)
    add_table(doc,
        ["Product", "Pattern for PINITHub"],
        [
            ["GitHub Enterprise", "Org settings left nav, audit log, team repos"],
            ["Notion", "Workspace switcher, nested hierarchy"],
            ["Slack", "Invites, @mentions, notification groups"],
            ["Microsoft 365", "Admin center, compliance center, role templates"],
            ["Dropbox Business", "Shared folders, viewer links, activity feed"],
            ["Atlassian", "Investigation = Jira issue workflow"],
        ],
    )

    doc.add_heading("Implementation Phasing (Recommended)", 1)
    add_table(doc,
        ["Phase", "Scope", "Validates"],
        [
            ["E1", "Org creation post-payment, Owner role, dashboard shell", "Purchase → setup journey"],
            ["E2", "Members, invites, roles", "Team collaboration"],
            ["E3", "Shared vault folders + ACL", "Core enterprise value"],
            ["E4", "Investigation cases + comments + assign", "Multi-user forensics"],
            ["E5", "Audit logs + compliance export", "Governance"],
            ["E6", "API keys, webhooks, billing seats", "Developer + finance"],
            ["E7", "Monitoring queue, scheduled reports", "Full enterprise loop"],
        ],
    )

    doc.add_heading("Summary", 1)
    doc.add_paragraph(
        "This specification defines PINITHub Enterprise as an organization-centric layer on top of the existing "
        "DNA → Vault → Investigation core loop. The experience mirrors best-in-class SaaS: guided post-purchase "
        "onboarding, governance-heavy sidebar, shared vault with folder ACLs, Jira-style investigations, "
        "immutable audit logs, and admin/billing/API surfaces scoped by role."
    )
    doc.add_paragraph(
        "Next step before code: Review with legal, security, and sales stakeholders; validate Phase E1–E3 for first Enterprise MVP."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
