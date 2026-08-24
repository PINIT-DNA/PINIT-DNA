"""Convert Database_Structure markdown to Word .docx (one-off utility)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def add_formatted_run(paragraph, text: str, bold=False, italic=False, mono=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if mono:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
    return run


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [strip_md_inline(c.strip()) for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < cols:
                table.rows[i].cells[j].text = cell
    doc.add_paragraph()


def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    # Title
    title = doc.add_heading("Database Structure — Pinit Hub, Exchange & Landing", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip().removeprefix("```").strip()
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                label = f"[Diagram: {code_lang}] " if code_lang == "mermaid" else ""
                add_formatted_run(p, label, bold=True)
                for cl in code_lines:
                    cp = doc.add_paragraph()
                    add_formatted_run(cp, cl, mono=True)
                doc.add_paragraph()
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(strip_md_inline(line[2:]), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(strip_md_inline(line[3:]), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(strip_md_inline(line[4:]), level=3)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(strip_md_inline(line[5:]), level=4)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table_rows(table_lines))
            continue

        if line.strip().startswith("*") and line.strip().endswith("*"):
            p = doc.add_paragraph(strip_md_inline(line.strip().strip("*")))
            p.runs[0].italic = True
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", line)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    add_formatted_run(p, part[2:-2], bold=True)
                elif part.startswith("`") and part.endswith("`"):
                    add_formatted_run(p, part[1:-1], mono=True)
                else:
                    add_formatted_run(p, part)
        i += 1

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "Database_Structure_Pinit_Hub.md"
    out = root / "docs" / "Database_Structure_Pinit_Hub.docx"
    md_to_docx(md, out)
