#!/usr/bin/env python3
"""Convert TEAM-ONBOARDING-PROJECT-GUIDE.md to a Word .docx file."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "TEAM-ONBOARDING-PROJECT-GUIDE.md"
OUT_PATH = ROOT / "docs" / "PinIT-Hub-Team-Onboarding-Guide.docx"


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_horizontal_line(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B4F72")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_formatted_runs(paragraph, text: str, base_size: int = 11):
    """Parse inline markdown: **bold**, `code`, [link](url)."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 1, color=RGBColor(0xC0, 0x39, 0x2B))
        elif token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(f"{label} ({url})")
                set_run_font(run, size=base_size, color=RGBColor(0x1A, 0x52, 0x7A))
            else:
                run = paragraph.add_run(token)
                set_run_font(run, size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            val = row[j] if j < len(row) else ""
            # strip markdown bold for cell clarity
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            run = p.add_run(clean)
            set_run_font(run, size=9, bold=(i == 0))
            if i == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1B4F72")
                shading.set(qn("w:val"), "clear")
                cell._teProp = cell._tc.get_or_add_tcPr()
                cell._tc.get_or_add_tcPr().append(shading)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # skip separator |---|---|
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def convert():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    i = 0
    in_code = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                # Add code block as monospace paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.15)
                code_text = "\n".join(code_lines)
                run = p.add_run(code_text if code_text else " ")
                set_run_font(run, name="Consolas", size=8, color=RGBColor(0x2C, 0x3E, 0x50))
                # light background via shading on paragraph
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F4F6F7")
                p._p.get_or_add_pPr().append(shd)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            add_horizontal_line(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # strip anchor links like {#...} if any
            title = re.sub(r"\{#.*\}$", "", title).strip()
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", title)

            if level == 1:
                p = doc.add_heading(title, level=0)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
            else:
                p = doc.add_heading(title, level=min(level, 3))
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, next_i = parse_table_block(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote = re.sub(r"^>\s?", "", stripped)
            # collect consecutive quote lines
            quotes = [quote]
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                quotes.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, " ".join(quotes), base_size=11)
            for run in p.runs:
                run.italic = True
                if run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
            continue

        # Unordered list
        if re.match(r"^[-*]\s+", stripped):
            item = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, item, base_size=11)
            i += 1
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            item = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, item, base_size=11)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped, base_size=11)
        i += 1

    # Footer note
    add_horizontal_line(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PinIT Hub — Confidential Team Documentation | Secure · Connect · Control")
    set_run_font(run, size=9, color=RGBColor(0x7F, 0x8C, 0x8D))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert()
