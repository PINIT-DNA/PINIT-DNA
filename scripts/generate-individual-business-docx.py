#!/usr/bin/env python3
"""Generate a black-and-white Word doc for PinIT Hub Individual/Business current flow."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "PinIT_Hub_Individual_Business_Current_Flow.md"
OUT_PATH = ROOT / "docs" / "PinIT_Hub_Individual_Business_Current_Flow.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = BLACK


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_horizontal_line(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_formatted_runs(paragraph, text: str, base_size: int = 11, bold_all: bool = False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, bold=bold_all)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=base_size - 1)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, bold=bold_all)


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, text, base_size=9, bold_all=(r_idx == 0))
            set_cell_border(cell)
            if r_idx == 0:
                shade_cell(cell, "E8E8E8")  # light gray header only — still B&W
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = BLACK


def convert():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level in range(1, 5):
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.color.rgb = BLACK
            hs.font.name = "Calibri"
        except KeyError:
            pass

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Consolas", size=9)
                p.paragraph_format.left_indent = Inches(0.15)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            add_horizontal_line(doc)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", title)
            p = doc.add_heading(title, level=min(level, 3) if level > 1 else 0)
            for run in p.runs:
                run.font.color.rgb = BLACK
                run.font.name = "Calibri"
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, next_i = parse_table_block(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, num.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    # Force all remaining runs black
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = BLACK

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("PinIT Hub — Individual & Business Current Flow · Confidential · Black & White")
    set_run_font(run, size=8)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert()
