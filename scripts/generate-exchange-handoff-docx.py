#!/usr/bin/env python3
"""Generate clean black-and-white Word handoff for Exchange builder."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "PinIT_Exchange_Builder_Handoff.md"
OUT_PATH = ROOT / "docs" / "PinIT_Exchange_Builder_Handoff.docx"
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = BLACK


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_line(doc: Document):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_runs(paragraph, text: str, size: int = 11, bold_all: bool = False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size, bold=bold_all)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(8, size - 1))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=bold_all)


def parse_table(lines: list[str], start: int):
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            add_runs(cell.paragraphs[0], row[c_idx] if c_idx < len(row) else "", size=9, bold_all=(r_idx == 0))
            set_cell_border(cell)
            if r_idx == 0:
                shade_cell(cell, "EEEEEE")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = BLACK


def convert():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    for level in range(1, 5):
        try:
            doc.styles[f"Heading {level}"].font.color.rgb = BLACK
        except KeyError:
            pass

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, name="Consolas", size=9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            add_line(doc)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = re.sub(r"\*\*([^*]+)\*\*", r"\1", heading.group(2).strip())
            p = doc.add_heading(title, level=0 if level == 1 else min(level, 3))
            for run in p.runs:
                run.font.color.rgb = BLACK
                run.font.name = "Calibri"
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, numbered.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.color.rgb = BLACK

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PinIT Exchange Builder Handoff · Share with Exchange teammate")
    set_run_font(run, size=8)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert()
